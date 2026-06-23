#!/usr/bin/env python3
import json, os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE   = os.path.dirname(os.path.abspath(__file__))
SHOTS  = os.path.join(BASE, 'capturas')
DATA   = os.path.join(BASE, 'resultados_pruebas.json')
OUTPUT = os.path.join(BASE, 'Informe_Pruebas_Funcionales_Mara_Beauty_Studio.docx')

with open(DATA) as f:
    data = json.load(f)

resultados = data['resultados']
shots      = data['shots']
fecha_raw  = data['fecha']
fecha      = datetime.fromisoformat(fecha_raw).strftime('%d/%m/%Y %H:%M')

# ── Colores ─────────────────────────────────────────────────────────
COLOR_TITULO   = RGBColor(0x1A, 0x1A, 0x2E)   # azul oscuro
COLOR_ACCENT   = RGBColor(0xC6, 0x7B, 0xA3)   # rosa/morado Mara
COLOR_CORRECTO = RGBColor(0x22, 0x8B, 0x22)   # verde
COLOR_ERROR    = RGBColor(0xCC, 0x00, 0x00)   # rojo
COLOR_WARN     = RGBColor(0xE6, 0x7E, 0x00)   # naranja
COLOR_INFO     = RGBColor(0x33, 0x66, 0xCC)   # azul
COLOR_BORDE    = RGBColor(0xE5, 0xE5, 0xE5)   # gris claro
COLOR_FILA     = RGBColor(0xFA, 0xF0, 0xF5)   # rosa pálido

def resultado_color(r):
    if r == 'CORRECTO':    return COLOR_CORRECTO
    if r == 'ERROR':       return COLOR_ERROR
    if r == 'ADVERTENCIA': return COLOR_WARN
    return COLOR_INFO

def resultado_icon(r):
    if r == 'CORRECTO':    return '✓'
    if r == 'ERROR':       return '✗'
    if r == 'ADVERTENCIA': return '⚠'
    return '●'

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)

def set_cell_border(cell, color='E5E5E5'):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = COLOR_TITULO
    elif level == 2:
        run.font.size = Pt(12)
        run.font.color.rgb = COLOR_ACCENT
    else:
        run.font.size = Pt(10)
        run.font.color.rgb = COLOR_TITULO
    return p

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run('─' * 72)
    run.font.size = Pt(6)
    run.font.color.rgb = COLOR_BORDE

def add_screenshot(doc, path, caption, width=15.5):
    if not os.path.exists(path):
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run()
    run.add_picture(path, width=Cm(width))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(10)
    cap_run = cap.runs[0]
    cap_run.font.size = Pt(8)
    cap_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    cap_run.italic = True

# ────────────────────────────────────────────────────────────────────
doc = Document()

# Márgenes ajustados
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# Fuente base
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)

# ══ PORTADA ════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
t1 = title_p.add_run('MARA BEAUTY STUDIO')
t1.font.size = Pt(26)
t1.font.bold = True
t1.font.color.rgb = COLOR_TITULO

doc.add_paragraph()
sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
t2 = sub_p.add_run('Informe de Pruebas Funcionales')
t2.font.size = Pt(16)
t2.font.color.rgb = COLOR_ACCENT
t2.bold = True

doc.add_paragraph()
info_p = doc.add_paragraph()
info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
t3 = info_p.add_run(f'Fecha: {fecha}  |  Entorno: Producción (Netlify + Cloudflare Tunnel)')
t3.font.size = Pt(9)
t3.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

doc.add_paragraph()
doc.add_paragraph()

# ── Resumen estadístico ─────────────────────────────────────────────
total    = len(resultados)
correctos = sum(1 for r in resultados if r['resultado'] == 'CORRECTO')
errores   = sum(1 for r in resultados if r['resultado'] == 'ERROR')
warnings  = sum(1 for r in resultados if r['resultado'] == 'ADVERTENCIA')
infos     = sum(1 for r in resultados if r['resultado'] == 'INFORMATIVO')

stats_table = doc.add_table(rows=1, cols=4)
stats_table.alignment = WD_TABLE_ALIGNMENT.CENTER
stats_table.style = 'Table Grid'
for i, (label, val, color) in enumerate([
    ('CORRECTAS', str(correctos), 'C6EFCE'),
    ('ERRORES',   str(errores),   'FFC7CE'),
    ('ADVERTENCIAS', str(warnings), 'FFEB9C'),
    ('TOTAL',     str(total),     'DDEEFF'),
]):
    cell = stats_table.rows[0].cells[i]
    set_cell_bg(cell, color)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    r1 = p.add_run(f'{val}\n')
    r1.font.size = Pt(22)
    r1.font.bold = True
    r2 = p.add_run(label)
    r2.font.size = Pt(8)
    r2.font.bold = True
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)

doc.add_page_break()

# ══ 1. VERIFICACIÓN DE DOMINIO ══════════════════════════════════════
add_heading(doc, '1. Verificación de Dominio', 1)
add_heading(doc, 'URL evaluada: http://genuine-raindrop-4a79b7.netlify.app', 3)

domain_tests = [r for r in resultados if any(k in r['prueba'] for k in ['DNS','HTTP','HTTPS','Ping','Whois'])]
if domain_tests:
    tbl = doc.add_table(rows=1 + len(domain_tests), cols=3)
    tbl.style = 'Table Grid'
    headers = ['Prueba', 'Resultado', 'Detalle']
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]
        set_cell_bg(c, '1A1A2E')
        p = c.paragraphs[0]
        r = p.add_run(h)
        r.font.bold  = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size  = Pt(9)
    for j, test in enumerate(domain_tests):
        row = tbl.rows[j+1]
        row.cells[0].text = test['prueba']
        ico = resultado_icon(test['resultado'])
        r2 = row.cells[1].paragraphs[0]
        run = r2.add_run(f"{ico} {test['resultado']}")
        run.font.color.rgb = resultado_color(test['resultado'])
        run.font.bold = True
        run.font.size = Pt(9)
        row.cells[2].text = test['obs']
        bg = 'FAF0F5' if j % 2 else 'FFFFFF'
        for ci in row.cells:
            set_cell_bg(ci, bg)
            ci.paragraphs[0].paragraph_format.space_before = Pt(3)
            ci.paragraphs[0].paragraph_format.space_after  = Pt(3)
            for run2 in ci.paragraphs[0].runs:
                run2.font.size = Pt(9)
    tbl.columns[0].width = Cm(6)
    tbl.columns[1].width = Cm(3.5)
    tbl.columns[2].width = Cm(8)

add_divider(doc)

# ══ 2. RESPONSIVE — CAPTURAS POR DISPOSITIVO ═════════════════════════
doc.add_page_break()
add_heading(doc, '2. Diseño Responsive — Múltiples Dispositivos', 1)

responsive_map = [
    ('responsive_Desktop_1920',   'Desktop — 1920 × 1080 px'),
    ('responsive_Laptop_1366',    'Laptop — 1366 × 768 px'),
    ('responsive_Tablet_iPad',    'Tablet iPad — 768 × 1024 px'),
    ('responsive_Mobile_iPhone',  'iPhone 14 — 390 × 844 px'),
    ('responsive_Mobile_Samsung', 'Samsung Galaxy — 360 × 800 px'),
]
for key, label in responsive_map:
    if key in shots and os.path.exists(shots[key]):
        add_heading(doc, label, 2)
        add_screenshot(doc, shots[key], f'Vista responsive: {label}', width=15)
        add_divider(doc)

# ══ 3. PRUEBAS MULTI-NAVEGADOR ══════════════════════════════════════
doc.add_page_break()
add_heading(doc, '3. Compatibilidad Multi-Navegador', 1)

nav_map = [
    ('nav_Chrome 120',   'Google Chrome 120'),
    ('nav_Firefox 121',  'Mozilla Firefox 121'),
    ('nav_Edge 120',     'Microsoft Edge 120'),
    ('nav_Safari 17',    'Apple Safari 17'),
]
for key, label in nav_map:
    if key in shots and os.path.exists(shots[key]):
        add_heading(doc, label, 2)
        add_screenshot(doc, shots[key], f'Renderizado en {label}', width=15)
        add_divider(doc)

# ══ 4. FUNCIONALIDADES PRINCIPALES ══════════════════════════════════
doc.add_page_break()
add_heading(doc, '4. Funcionalidades Principales', 1)

func_map = [
    ('inicio',            '4.1 Página Principal — Calendario de citas'),
    ('nav_siguiente',     '4.2 Navegación — Semana siguiente'),
    ('nav_anterior',      '4.3 Navegación — Semana anterior'),
    ('sin_slots_disponibles', '4.4 Calendario sin slots disponibles'),
    ('admin_login',       '4.5 Módulo Admin — Formulario de Login'),
    ('login_vacio',       '4.6 Validación — Campos vacíos'),
    ('login_invalido',    '4.7 Validación — Credenciales incorrectas'),
    ('login_exitoso',     '4.8 Login exitoso — Redirección al Dashboard'),
    ('dashboard',         '4.9 Dashboard Administrativo'),
    ('crear_cita',        '4.10 Crear Cita — Modal de creación'),
    ('simulador',         '4.11 Simulador de Correos'),
    ('404',               '4.12 Manejo de Páginas No Encontradas'),
]
for key, label in func_map:
    if key in shots and os.path.exists(shots[key]):
        add_heading(doc, label, 2)
        add_screenshot(doc, shots[key], label, width=15)
        add_divider(doc)

# ══ 5. TABLA RESUMEN COMPLETA ════════════════════════════════════════
doc.add_page_break()
add_heading(doc, '5. Tabla de Resultados', 1)

tbl2 = doc.add_table(rows=1 + len(resultados), cols=3)
tbl2.style = 'Table Grid'
for i, h in enumerate(['Prueba', 'Resultado', 'Observación']):
    c = tbl2.rows[0].cells[i]
    set_cell_bg(c, '1A1A2E')
    p = c.paragraphs[0]
    rh = p.add_run(h)
    rh.font.bold  = True
    rh.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    rh.font.size  = Pt(9)

for j, test in enumerate(resultados):
    row = tbl2.rows[j+1]
    row.cells[0].text = test['prueba']
    p2 = row.cells[1].paragraphs[0]
    ico = resultado_icon(test['resultado'])
    run2 = p2.add_run(f"{ico} {test['resultado']}")
    run2.font.color.rgb = resultado_color(test['resultado'])
    run2.font.bold = True
    run2.font.size = Pt(8)
    row.cells[2].text = test['obs']
    bg = 'FAF0F5' if j % 2 else 'FFFFFF'
    for ci in row.cells:
        set_cell_bg(ci, bg)
        ci.paragraphs[0].paragraph_format.space_before = Pt(2)
        ci.paragraphs[0].paragraph_format.space_after  = Pt(2)
        for r3 in ci.paragraphs[0].runs:
            r3.font.size = Pt(8)

tbl2.columns[0].width = Cm(7)
tbl2.columns[1].width = Cm(3)
tbl2.columns[2].width = Cm(7)

# ══ PIE DE PÁGINA ════════════════════════════════════════════════════
doc.add_paragraph()
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = footer_p.add_run(f'Mara Beauty Studio — Informe generado automáticamente el {fecha}')
fr.font.size = Pt(8)
fr.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
fr.italic = True

doc.save(OUTPUT)
print(f'✅ Informe generado: {OUTPUT}')
