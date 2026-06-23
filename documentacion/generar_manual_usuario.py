#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os

SHOTS = '/Users/andres/Documents/Mara Beauty Studio/capturas'
OUT   = '/Users/andres/Documents/Mara Beauty Studio/Manual_Usuario_Mara_Beauty_Studio.docx'
# ── Numeración de secciones actualizada ──
# 1. Objetivos del Software   ← NUEVO
# 2. Presentación del Sistema
# 3. Descripción de la Interfaz
# 4. Instrucciones de Ingreso
# 5. Módulo Cliente
# 6. Módulo Administrador
# 7. Flujos de Trabajo
# 8. Compatibilidad y Dispositivos
# 9. Preguntas Frecuentes

# ── Paleta ────────────────────────────────────────────────────────────
C_DARK  = RGBColor(0x1A, 0x1A, 0x2E)
C_ROSE  = RGBColor(0xC6, 0x7B, 0xA3)
C_GRAY  = RGBColor(0x88, 0x88, 0x88)
C_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
C_GREEN = RGBColor(0x2E, 0x7D, 0x32)
C_BLUE  = RGBColor(0x15, 0x65, 0xC0)

def shd(cell, hex_col):
    tc = cell._tc; pr = tc.get_or_add_tcPr()
    s = OxmlElement('w:shd')
    s.set(qn('w:fill'), hex_col); s.set(qn('w:val'), 'clear')
    pr.append(s)

def border_cell(cell, col='DDDDDD'):
    tc = cell._tc; pr = tc.get_or_add_tcPr()
    tb = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '4'); b.set(qn('w:color'), col)
        tb.append(b)
    pr.append(tb)

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.2); s.bottom_margin = Cm(2.2)
    s.left_margin = Cm(2.8); s.right_margin = Cm(2.8)
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(10)

# ── Helpers ───────────────────────────────────────────────────────────
def H1(texto):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(texto.upper())
    r.bold = True; r.font.size = Pt(14); r.font.color.rgb = C_DARK
    line = doc.add_paragraph()
    line.paragraph_format.space_before = Pt(0)
    line.paragraph_format.space_after  = Pt(10)
    lr = line.add_run('─' * 78)
    lr.font.size = Pt(5); lr.font.color.rgb = C_ROSE

def H2(texto):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(f'◆  {texto}')
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = C_ROSE

def H3(texto):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(f'▸  {texto}')
    r.bold = True; r.font.size = Pt(10); r.font.color.rgb = C_DARK

def body(texto):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(texto)
    r.font.size = Pt(10); r.font.color.rgb = C_DARK

def nota(texto, tipo='info'):
    colores = {'info': 'DDEEFF', 'tip': 'E8F5E9', 'warn': 'FFF8E1'}
    iconos  = {'info': 'ℹ', 'tip': '✔', 'warn': '⚠'}
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    c = tbl.rows[0].cells[0]
    shd(c, colores.get(tipo, 'DDEEFF'))
    border_cell(c, 'AAAAAA')
    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(f'{iconos.get(tipo, "ℹ")}  {texto}')
    r.font.size = Pt(9)
    r.font.color.rgb = C_DARK
    tbl.columns[0].width = Cm(14.4)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def bullet(texto, negrita=''):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    if negrita:
        rb = p.add_run(negrita)
        rb.bold = True; rb.font.size = Pt(10)
    r = p.add_run(texto)
    r.font.size = Pt(10)

def paso(num, texto, detalle=''):
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    cn = tbl.rows[0].cells[0]
    ct = tbl.rows[0].cells[1]
    shd(cn, 'C67BA3'); border_cell(cn)
    shd(ct, 'FAF0F5'); border_cell(ct)
    cn.width = Cm(1.2); ct.width = Cm(13.2)
    pn = cn.paragraphs[0]
    pn.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pn.paragraph_format.space_before = Pt(4)
    rn = pn.add_run(str(num))
    rn.bold = True; rn.font.size = Pt(12); rn.font.color.rgb = C_WHITE
    pt = ct.paragraphs[0]
    pt.paragraph_format.space_before = Pt(4)
    pt.paragraph_format.space_after  = Pt(4)
    pt.paragraph_format.left_indent  = Cm(0.3)
    rb = pt.add_run(texto)
    rb.bold = True; rb.font.size = Pt(10)
    if detalle:
        rd = pt.add_run(f'\n{detalle}')
        rd.font.size = Pt(9); rd.font.color.rgb = C_GRAY
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def img(nombre_archivo, caption, ancho=14.5):
    ruta = os.path.join(SHOTS, nombre_archivo)
    if not os.path.exists(ruta):
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.add_run().add_picture(ruta, width=Cm(ancho))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(12)
    cr = cap.runs[0]
    cr.font.size = Pt(8); cr.font.color.rgb = C_GRAY; cr.italic = True

def tabla_simple(headers, filas, anchos=None):
    tbl = doc.add_table(rows=1+len(filas), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]
        shd(c, '1A1A2E'); border_cell(c)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        r = p.add_run(h)
        r.bold = True; r.font.size = Pt(9); r.font.color.rgb = C_WHITE
    for ri, fila in enumerate(filas):
        bg = 'FAF0F5' if ri % 2 else 'FFFFFF'
        for ci, val in enumerate(fila):
            c = tbl.rows[ri+1].cells[ci]
            shd(c, bg); border_cell(c)
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after  = Pt(3)
            p.paragraph_format.left_indent  = Cm(0.2)
            r = p.add_run(str(val))
            r.font.size = Pt(9)
    if anchos:
        for i, w in enumerate(anchos):
            for row in tbl.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def separador():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run('· · · · · · · · · · · · · · · · · · · · · · · · · · · ·')
    r.font.size = Pt(8); r.font.color.rgb = RGBColor(0xDD, 0xBB, 0xCC)

# ═══════════════════════════════════════════════════════════════════
#  PORTADA
# ═══════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('MARA BEAUTY STUDIO')
r.font.size = Pt(30); r.bold = True; r.font.color.rgb = C_DARK

doc.add_paragraph()
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('Manual de Usuario')
r2.font.size = Pt(20); r2.bold = True; r2.font.color.rgb = C_ROSE

doc.add_paragraph()
for linea, sz, color in [
    ('Sistema de Gestión de Citas', 13, C_DARK),
    ('Guía completa para Clientes y Administradores', 11, C_GRAY),
    (f'Versión 1.0  ·  {datetime.now().strftime("%B %Y")}', 10, C_GRAY),
]:
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(linea)
    r3.font.size = Pt(sz); r3.font.color.rgb = color

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
#  ÍNDICE
# ═══════════════════════════════════════════════════════════════════
H1('Tabla de Contenido')
toc = [
    ('1.', 'Objetivos del Software'),
    ('2.', 'Presentación del Sistema'),
    ('3.', 'Descripción de la Interfaz'),
    ('4.', 'Instrucciones de Ingreso al Sistema'),
    ('5.', 'Módulo Cliente — Reserva y Cancelación de Citas'),
    ('6.', 'Módulo Administrador — Gestión del Sistema'),
    ('7.', 'Flujos de Trabajo'),
    ('8.', 'Compatibilidad y Dispositivos'),
    ('9.', 'Preguntas Frecuentes'),
]
for num, titulo in toc:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(f'   {num}   {titulo}')
    r.font.size = Pt(11); r.font.color.rgb = C_DARK

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
#  1. OBJETIVOS DEL SOFTWARE
# ═══════════════════════════════════════════════════════════════════
H1('1. Objetivos del Software')

H2('1.1 Objetivo general')
# Caja destacada para el objetivo general
tbl_obj = doc.add_table(rows=1, cols=1)
tbl_obj.style = 'Table Grid'
c_obj = tbl_obj.rows[0].cells[0]
shd(c_obj, 'F3E6EF'); border_cell(c_obj, 'C67BA3')
p_obj = c_obj.paragraphs[0]
p_obj.paragraph_format.space_before = Pt(8)
p_obj.paragraph_format.space_after  = Pt(8)
p_obj.paragraph_format.left_indent  = Cm(0.4)
r_obj = p_obj.add_run(
    'Desarrollar una aplicación web que permita a los clientes del salón Mara Beauty Studio '
    'consultar la disponibilidad de citas y realizar reservas en línea de forma autónoma, '
    'y al administrador gestionar eficientemente la agenda semanal del negocio, '
    'optimizando el proceso de agendamiento y reduciendo la gestión manual de citas.'
)
r_obj.font.size = Pt(10); r_obj.font.color.rgb = C_DARK
tbl_obj.columns[0].width = Cm(14.4)
doc.add_paragraph().paragraph_format.space_after = Pt(6)

H2('1.2 Objetivos específicos')
objetivos = [
    ('Facilitar el acceso a la agenda',
     'Permitir que cualquier cliente pueda consultar los horarios disponibles del salón '
     'sin necesidad de llamar por teléfono o desplazarse físicamente.'),
    ('Automatizar el proceso de reserva',
     'Brindar al cliente la posibilidad de agendar su cita de forma autónoma a cualquier '
     'hora del día, ingresando únicamente su nombre y correo electrónico.'),
    ('Centralizar la gestión de la agenda',
     'Ofrecer al administrador un panel de control desde donde pueda crear, visualizar '
     'y eliminar los slots de citas de la semana de manera rápida e intuitiva.'),
    ('Notificar automáticamente a los clientes',
     'Enviar correos electrónicos de confirmación y cancelación de cita al momento en '
     'que se registra la acción, manteniendo al cliente informado en todo momento.'),
    ('Garantizar la disponibilidad en tiempo real',
     'Evitar dobles reservas o conflictos de horario mediante validaciones automáticas '
     'que aseguran que cada slot solo pueda ser reservado por un cliente a la vez.'),
    ('Ofrecer acceso multiplataforma',
     'Diseñar la interfaz con tecnología responsive para que el sistema funcione '
     'correctamente en computadores, tabletas y celulares sin instalar ninguna aplicación.'),
]
for titulo, descripcion in objetivos:
    tbl_e = doc.add_table(rows=1, cols=2)
    tbl_e.style = 'Table Grid'
    c_icon = tbl_e.rows[0].cells[0]
    c_txt  = tbl_e.rows[0].cells[1]
    shd(c_icon, 'C67BA3'); border_cell(c_icon, 'C67BA3')
    shd(c_txt,  'FDFBFD'); border_cell(c_txt,  'DDDDDD')
    c_icon.width = Cm(0.8); c_txt.width = Cm(13.6)
    pi = c_icon.paragraphs[0]
    pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pi.paragraph_format.space_before = Pt(6)
    ri = pi.add_run('✦')
    ri.font.size = Pt(11); ri.font.color.rgb = C_WHITE; ri.bold = True
    pt2 = c_txt.paragraphs[0]
    pt2.paragraph_format.space_before = Pt(5)
    pt2.paragraph_format.space_after  = Pt(5)
    pt2.paragraph_format.left_indent  = Cm(0.3)
    rb2 = pt2.add_run(f'{titulo}:  ')
    rb2.bold = True; rb2.font.size = Pt(10)
    rd2 = pt2.add_run(descripcion)
    rd2.font.size = Pt(10); rd2.font.color.rgb = C_DARK
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

separador()

H2('1.3 ¿Para qué sirve el sistema?')
tabla_simple(
    ['Necesidad del negocio', 'Solución que ofrece el sistema'],
    [
        ['Los clientes llaman por teléfono para saber horarios disponibles',
         'El cliente consulta el calendario en línea 24/7 desde cualquier dispositivo'],
        ['El administrador lleva la agenda en papel o mensajes de WhatsApp',
         'Toda la agenda queda centralizada en el sistema, accesible desde el navegador'],
        ['Es difícil confirmar o recordar citas a los clientes',
         'El sistema envía correos automáticos de confirmación y cancelación'],
        ['Riesgo de agendar dos clientes en el mismo horario',
         'El sistema bloquea el slot en tiempo real al momento de la reserva'],
        ['El cliente no sabe cómo cancelar si no puede asistir',
         'El correo de confirmación incluye un enlace directo para cancelar la cita'],
    ],
    anchos=[7, 9]
)

separador()

H2('1.4 ¿Quiénes utilizan el sistema?')

# Tarjeta Cliente
tbl_u1 = doc.add_table(rows=1, cols=2)
tbl_u1.style = 'Table Grid'
ch1 = tbl_u1.rows[0].cells[0]; ct1 = tbl_u1.rows[0].cells[1]
shd(ch1, '1A1A2E'); border_cell(ch1); shd(ct1, 'FAF4F8'); border_cell(ct1)
ch1.width = Cm(3.5); ct1.width = Cm(10.9)
ph1 = ch1.paragraphs[0]
ph1.alignment = WD_ALIGN_PARAGRAPH.CENTER
ph1.paragraph_format.space_before = Pt(10)
ph1.paragraph_format.space_after  = Pt(10)
rh1a = ph1.add_run('👤\n')
rh1a.font.size = Pt(18)
rh1b = ph1.add_run('CLIENTE')
rh1b.bold = True; rh1b.font.size = Pt(11); rh1b.font.color.rgb = C_WHITE
pct1 = ct1.paragraphs[0]
pct1.paragraph_format.space_before = Pt(8)
pct1.paragraph_format.left_indent  = Cm(0.4)
pct1.add_run('Perfil:\n').bold = True
pct1.runs[-1].font.size = Pt(10)
for item in [
    'Persona que desea agendar un servicio en el salón',
    'No requiere conocimientos técnicos ni registro previo',
    'Accede desde su celular, tableta o computador',
    'Solo necesita su nombre y correo para reservar',
]:
    p_li = ct1.add_paragraph()
    p_li.paragraph_format.left_indent = Cm(0.6)
    p_li.paragraph_format.space_after = Pt(1)
    r_li = p_li.add_run(f'•  {item}')
    r_li.font.size = Pt(9); r_li.font.color.rgb = C_DARK
p_acc = ct1.add_paragraph()
p_acc.paragraph_format.left_indent = Cm(0.4)
p_acc.paragraph_format.space_before = Pt(4)
p_acc.paragraph_format.space_after  = Pt(8)
ra = p_acc.add_run('Acceso:  ')
ra.bold = True; ra.font.size = Pt(9)
rb_acc = p_acc.add_run('Público — sin contraseña')
rb_acc.font.size = Pt(9); rb_acc.font.color.rgb = C_GREEN
tbl_u1.columns[0].width = Cm(3.5)
tbl_u1.columns[1].width = Cm(10.9)
doc.add_paragraph().paragraph_format.space_after = Pt(6)

# Tarjeta Administrador
tbl_u2 = doc.add_table(rows=1, cols=2)
tbl_u2.style = 'Table Grid'
ch2 = tbl_u2.rows[0].cells[0]; ct2 = tbl_u2.rows[0].cells[1]
shd(ch2, 'C67BA3'); border_cell(ch2); shd(ct2, 'FDF6FB'); border_cell(ct2)
ch2.width = Cm(3.5); ct2.width = Cm(10.9)
ph2 = ch2.paragraphs[0]
ph2.alignment = WD_ALIGN_PARAGRAPH.CENTER
ph2.paragraph_format.space_before = Pt(10)
ph2.paragraph_format.space_after  = Pt(10)
rh2a = ph2.add_run('🔑\n')
rh2a.font.size = Pt(18)
rh2b = ph2.add_run('ADMINISTRADOR')
rh2b.bold = True; rh2b.font.size = Pt(10); rh2b.font.color.rgb = C_WHITE
pct2 = ct2.paragraphs[0]
pct2.paragraph_format.space_before = Pt(8)
pct2.paragraph_format.left_indent  = Cm(0.4)
pct2.add_run('Perfil:\n').bold = True
pct2.runs[-1].font.size = Pt(10)
for item in [
    'Dueño o encargado del salón Mara Beauty Studio',
    'Gestiona la agenda semanal de citas del negocio',
    'Tiene acceso exclusivo al panel de administración',
    'Puede crear y eliminar slots, y ver todos los datos de clientes',
]:
    p_li2 = ct2.add_paragraph()
    p_li2.paragraph_format.left_indent = Cm(0.6)
    p_li2.paragraph_format.space_after = Pt(1)
    r_li2 = p_li2.add_run(f'•  {item}')
    r_li2.font.size = Pt(9); r_li2.font.color.rgb = C_DARK
p_acc2 = ct2.add_paragraph()
p_acc2.paragraph_format.left_indent = Cm(0.4)
p_acc2.paragraph_format.space_before = Pt(4)
p_acc2.paragraph_format.space_after  = Pt(8)
ra2 = p_acc2.add_run('Acceso:  ')
ra2.bold = True; ra2.font.size = Pt(9)
rb_acc2 = p_acc2.add_run('Privado — requiere usuario y contraseña')
rb_acc2.font.size = Pt(9); rb_acc2.font.color.rgb = C_ROSE
tbl_u2.columns[0].width = Cm(3.5)
tbl_u2.columns[1].width = Cm(10.9)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
#  2. PRESENTACIÓN
# ═══════════════════════════════════════════════════════════════════
H1('2. Presentación del Sistema')

body('Mara Beauty Studio es una aplicación web diseñada para gestionar las citas de un salón de belleza. '
     'Permite a los clientes consultar la disponibilidad y reservar citas en línea, y al administrador '
     'organizar los horarios, visualizar estadísticas y gestionar la agenda semanal.')

doc.add_paragraph()

H2('1.1 Usuarios del sistema')
tabla_simple(
    ['Tipo de usuario', 'Acceso', 'Funciones principales'],
    [
        ['Cliente', 'Público — sin contraseña', 'Ver calendario, reservar cita, cancelar cita desde email'],
        ['Administrador', 'Privado — usuario y contraseña', 'Crear/eliminar slots, ver agenda completa, estadísticas y clientes'],
    ],
    anchos=[3.5, 4, 8.5]
)

H2('1.2 Cómo acceder al sistema')
body('El sistema funciona directamente desde el navegador web. No requiere instalar ninguna aplicación.')
bullet('Dirección web: ', 'http://genuine-raindrop-4a79b7.netlify.app')
bullet('Navegadores compatibles: ', 'Google Chrome, Mozilla Firefox, Microsoft Edge, Safari')
bullet('Dispositivos: ', 'Computador, tableta y celular (diseño adaptable)')

nota('No necesita crear una cuenta para reservar una cita. Solo ingrese su nombre y correo electrónico al momento de reservar.', 'tip')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
#  2. DESCRIPCIÓN DE LA INTERFAZ
# ═══════════════════════════════════════════════════════════════════
H1('2. Descripción de la Interfaz')

H2('2.1 Página principal — Vista del Cliente')
body('Al ingresar al sitio, el cliente encuentra directamente el calendario semanal de citas. '
     'No se requiere ningún registro previo.')

img('02_pagina_principal.png', 'Figura 1 — Página principal con el calendario semanal de citas')

H3('Componentes de la página principal')
tabla_simple(
    ['Elemento', 'Descripción'],
    [
        ['Encabezado',         'Nombre del salón y acceso al área de administración (botón "Iniciar sesión")'],
        ['Título "Calendario"','Indica la sección de visualización de citas disponibles'],
        ['Botones < >',        'Permiten navegar entre semanas (anterior y siguiente)'],
        ['Indicador de semana','Muestra la semana y año que se está visualizando actualmente'],
        ['Grilla de citas',    'Columnas por día (Lun–Dom) y filas por hora (08:00–18:00)'],
        ['Slot ABIERTO',       'Casilla verde — cita disponible para reservar'],
        ['Slot RESERVADA',     'Casilla de otro color — cita ya ocupada, no se puede seleccionar'],
        ['Botón "Reservar cita"', 'Se activa al seleccionar un slot disponible'],
    ],
    anchos=[4, 12]
)

separador()

H2('2.2 Panel de Administración — Dashboard')
body('El administrador accede a un panel exclusivo después de iniciar sesión. Desde aquí gestiona '
     'toda la agenda del salón.')

img('11_dashboard_admin.png', 'Figura 2 — Dashboard del administrador con la agenda semanal')

H3('Componentes del dashboard')
tabla_simple(
    ['Elemento', 'Descripción'],
    [
        ['Calendario admin',      'Misma grilla semanal, pero muestra el nombre del cliente en citas reservadas'],
        ['Botón "Agregar citas"', 'Abre el formulario para crear nuevos slots de citas disponibles'],
        ['Botón "Eliminar"',      'Elimina el slot seleccionado (solo si está disponible, no reservado)'],
        ['Botón "Cerrar sesión"', 'Cierra la sesión y regresa al formulario de login'],
        ['Navegación < >',        'Igual que en la vista del cliente — navega entre semanas'],
    ],
    anchos=[4, 12]
)

separador()

H2('2.3 Formulario de Login — Área de Administrador')
body('Página de acceso exclusivo para el administrador del salón. Se accede desde el botón '
     '"Iniciar sesión" en la esquina superior de la página principal.')

img('07_admin_login.png', 'Figura 3 — Formulario de ingreso al área de administración')

H3('Componentes del formulario')
tabla_simple(
    ['Campo / Botón', 'Descripción'],
    [
        ['Campo "Nombre de usuario"', 'Ingrese el nombre de usuario asignado al administrador'],
        ['Campo "Contraseña"',        'Ingrese la contraseña. Los caracteres se ocultan por seguridad'],
        ['Botón "Iniciar sesión"',    'Envía las credenciales y, si son correctas, accede al dashboard'],
        ['Mensaje de error',          'Aparece en rojo si el usuario o contraseña son incorrectos'],
    ],
    anchos=[5, 11]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
#  3. INSTRUCCIONES DE INGRESO
# ═══════════════════════════════════════════════════════════════════
H1('3. Instrucciones de Ingreso al Sistema')

H2('3.1 Ingreso como Cliente (sin contraseña)')
body('El cliente no necesita registrarse ni iniciar sesión. El acceso es inmediato.')

paso(1, 'Abrir el navegador web',
     'Chrome, Firefox, Edge o Safari en computador, tableta o celular.')
paso(2, 'Ingresar la dirección del sitio',
     'Escribir en la barra de direcciones: genuine-raindrop-4a79b7.netlify.app')
paso(3, 'El sistema carga automáticamente',
     'Se muestra el calendario con la semana actual y los slots disponibles.')

nota('Si no ve citas disponibles en la semana actual, use el botón > para avanzar a la semana siguiente.', 'info')

separador()

H2('3.2 Ingreso como Administrador')
body('El administrador debe autenticarse con usuario y contraseña para acceder al panel de gestión.')

paso(1, 'Desde la página principal, hacer clic en "Iniciar sesión"',
     'Botón ubicado en la esquina superior derecha del encabezado.')
paso(2, 'Ingresar las credenciales',
     'Escribir el nombre de usuario y la contraseña en los campos correspondientes.')
paso(3, 'Hacer clic en "Iniciar sesión"',
     'Si las credenciales son correctas, el sistema redirige automáticamente al dashboard.')
paso(4, 'Acceso al Panel de Administración',
     'Aparece la vista del administrador con el calendario completo y las opciones de gestión.')

img('10_login_exitoso.png', 'Figura 4 — Redirección exitosa al dashboard tras el login')

H3('¿Qué pasa si las credenciales son incorrectas?')
body('El sistema muestra el mensaje "Usuario o contraseña incorrectos" en color rojo debajo del formulario. '
     'La página permanece en el login para que el usuario pueda intentarlo nuevamente.')

img('09_login_credenciales_invalidas.png', 'Figura 5 — Mensaje de error por credenciales inválidas')

nota('Por seguridad, si dejó algún campo vacío, el sistema no permitirá enviar el formulario.', 'warn')

img('08_login_campos_vacios.png', 'Figura 6 — Validación de campos vacíos en el formulario de login')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
#  4. MÓDULO CLIENTE
# ═══════════════════════════════════════════════════════════════════
H1('4. Módulo Cliente — Reserva y Cancelación de Citas')

H2('4.1 Consultar citas disponibles')
body('Al ingresar al sistema, el calendario muestra automáticamente la semana actual con todos '
     'los horarios configurados por el administrador.')

H3('Cómo interpretar el calendario')
tabla_simple(
    ['Color / Estado', 'Significado', '¿Se puede seleccionar?'],
    [
        ['Verde — ABIERTO',      'El horario está disponible para reservar',           'Sí — hacer clic para seleccionar'],
        ['Gris — RESERVADA',     'El horario ya fue tomado por otro cliente',           'No'],
        ['Casilla vacía',        'El administrador no habilitó ese horario',            'No'],
        ['Seleccionado (borde)', 'Slot elegido por el cliente para reservar',           'Ya seleccionado'],
    ],
    anchos=[4, 6.5, 5.5]
)

H2('4.2 Navegar entre semanas')
body('Use los botones de navegación para ver citas en otras semanas.')

paso(1, 'Botón  >  — Avanzar a la semana siguiente',
     'Muestra los slots disponibles en los próximos días.')
paso(2, 'Botón  <  — Regresar a la semana anterior',
     'Permite revisar semanas pasadas (solo informativo).')
paso(3, 'Indicador central',
     'Muestra el número de semana y el año que está visualizando.')

img('03_semana_siguiente.png', 'Figura 7 — Navegación a la semana siguiente')

separador()

H2('4.3 Reservar una cita')
body('Proceso completo para agendar una cita en el salón.')

paso(1, 'Seleccionar un slot disponible (verde)',
     'Hacer clic sobre el horario deseado. El slot quedará resaltado.')
paso(2, 'Hacer clic en el botón "Reservar cita"',
     'Se activa automáticamente al seleccionar un slot disponible.')
paso(3, 'Completar el formulario de reserva',
     'Ingresar su nombre completo y correo electrónico en los campos.')
paso(4, 'Hacer clic en "Confirmar"',
     'El sistema registra la reserva y envía un correo de confirmación a su email.')
paso(5, 'Confirmación en pantalla',
     'Aparece el mensaje "¡Cita reservada con éxito!" y el slot cambia a estado RESERVADA.')

nota('Recibirá un correo electrónico de confirmación con los detalles de su cita y un enlace para cancelarla si lo necesita.', 'tip')

separador()

H2('4.4 Cancelar una cita')
body('Existen dos formas de cancelar una cita ya reservada.')

H3('Opción A — Desde el correo electrónico (recomendada)')
paso(1, 'Abrir el correo de confirmación recibido tras la reserva', '')
paso(2, 'Hacer clic en el enlace "Cancelar Cita"',
     'El enlace lo lleva directamente al sitio con el modal de cancelación abierto.')
paso(3, 'Confirmar la cancelación',
     'Hacer clic en "Confirmar" en el modal que aparece.')
paso(4, 'El sistema libera el horario',
     'El slot vuelve a quedar disponible y recibirá un correo de cancelación.')

H3('Opción B — Desde el calendario')
paso(1, 'Ingresar al sitio web', '')
paso(2, 'Ubicar la cita reservada en el calendario', 'Navegar a la semana correspondiente.')
paso(3, 'Seleccionar el slot reservado y usar la opción de cancelar', '')

nota('Solo puede cancelar citas propias. Las citas pasadas no pueden cancelarse.', 'warn')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
#  5. MÓDULO ADMINISTRADOR
# ═══════════════════════════════════════════════════════════════════
H1('5. Módulo Administrador — Gestión del Sistema')

nota('Todas las funciones de esta sección requieren haber iniciado sesión como administrador.', 'warn')

H2('5.1 Crear slots de citas')
body('El administrador define los horarios disponibles para que los clientes puedan reservar. '
     'Los slots se crean por rangos de hora y se generan automáticamente cada hora.')

paso(1, 'Desde el dashboard, hacer clic en "Agregar citas"',
     'Se abre el formulario de creación de slots.')
paso(2, 'Seleccionar la fecha',
     'Elegir el día para el cual se desean habilitar citas.')
paso(3, 'Seleccionar hora de inicio y hora de fin',
     'Por ejemplo: 09:00 a 13:00 creará 4 slots: 09:00, 10:00, 11:00, 12:00.')
paso(4, 'Hacer clic en "Guardar"',
     'Los slots aparecen inmediatamente en el calendario como horarios disponibles (verde).')

img('12_crear_cita_modal.png', 'Figura 8 — Modal para crear nuevos slots de citas')

nota('No es posible crear dos citas en el mismo horario. El sistema detecta conflictos automáticamente.', 'info')

separador()

H2('5.2 Eliminar un slot de cita')
body('El administrador puede eliminar slots que hayan quedado disponibles y ya no se necesiten.')

paso(1, 'En el calendario, hacer clic sobre el slot disponible (verde) que desea eliminar',
     'El slot queda seleccionado.')
paso(2, 'Hacer clic en el botón "Eliminar"',
     'Se abre un modal de confirmación.')
paso(3, 'Confirmar la eliminación',
     'El slot desaparece del calendario.')

nota('No es posible eliminar citas que ya están RESERVADAS por un cliente. Primero debe cancelarse la reserva.', 'warn')

separador()

H2('5.3 Ver la agenda completa')
body('El dashboard muestra el calendario con información adicional respecto a la vista del cliente.')

tabla_simple(
    ['Lo que ve el Administrador', 'Lo que ve el Cliente'],
    [
        ['Nombre del cliente en slots reservados',    'Solo ve "RESERVADA" sin datos del cliente'],
        ['Slots disponibles para habilitar/eliminar', 'Slots disponibles para reservar'],
        ['Historial de semanas pasadas',              'Igual acceso a semanas anteriores'],
        ['Botones de gestión (Agregar / Eliminar)',   'Solo botón de reservar'],
    ],
    anchos=[8, 8]
)

img('14_dashboard_completo.png', 'Figura 9 — Vista completa del dashboard administrativo')

separador()

H2('5.4 Consultar información de clientes')
body('El sistema registra automáticamente los datos de cada cliente al momento de su primera reserva. '
     'El administrador puede consultar esta información desde el panel.')

H3('Datos almacenados por cliente')
bullet('Nombre completo')
bullet('Correo electrónico')
bullet('Historial de citas reservadas')
bullet('Fecha de primer registro en el sistema')

separador()

H2('5.5 Cerrar sesión')
paso(1, 'Hacer clic en el botón "Cerrar sesión"',
     'Ubicado en la parte superior del dashboard.')
paso(2, 'El sistema cierra la sesión de forma segura',
     'El token de autenticación se elimina y se redirige al formulario de login.')

nota('Cierre siempre la sesión al terminar, especialmente si usa un equipo compartido.', 'warn')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
#  6. FLUJOS DE TRABAJO
# ═══════════════════════════════════════════════════════════════════
H1('6. Flujos de Trabajo')

H2('6.1 Flujo — Cliente reserva una cita')
tabla_simple(
    ['Paso', 'Acción', 'Actor', 'Resultado'],
    [
        ['1', 'Ingresa al sitio web',                        'Cliente',        'Se carga el calendario semanal'],
        ['2', 'Navega hasta encontrar un horario disponible', 'Cliente',        'Ve slots verdes (ABIERTO)'],
        ['3', 'Hace clic en el slot deseado',                'Cliente',        'Slot queda seleccionado'],
        ['4', 'Hace clic en "Reservar cita"',                'Cliente',        'Se abre el formulario de datos'],
        ['5', 'Ingresa nombre y correo — confirma',          'Cliente',        'Cita registrada en el sistema'],
        ['6', 'Sistema actualiza el calendario',             'Sistema',        'Slot cambia a RESERVADA'],
        ['7', 'Sistema envía email de confirmación',         'Sistema / Email', 'Cliente recibe confirmación'],
    ],
    anchos=[1.2, 5.8, 2.8, 6.2]
)

H2('6.2 Flujo — Cliente cancela una cita')
tabla_simple(
    ['Paso', 'Acción', 'Actor', 'Resultado'],
    [
        ['1', 'Abre el email de confirmación',          'Cliente',        'Encuentra el enlace de cancelación'],
        ['2', 'Hace clic en "Cancelar Cita"',           'Cliente',        'Se abre el sitio con modal de cancelación'],
        ['3', 'Confirma la cancelación',                'Cliente',        'Cita cancelada en el sistema'],
        ['4', 'Sistema libera el horario',              'Sistema',        'Slot vuelve a estado DISPONIBLE'],
        ['5', 'Sistema envía email de cancelación',     'Sistema / Email', 'Cliente recibe notificación'],
    ],
    anchos=[1.2, 5.8, 2.8, 6.2]
)

H2('6.3 Flujo — Administrador habilita citas para una semana')
tabla_simple(
    ['Paso', 'Acción', 'Actor', 'Resultado'],
    [
        ['1', 'Inicia sesión en el sistema',                     'Admin',   'Accede al dashboard'],
        ['2', 'Navega a la semana que quiere configurar',         'Admin',   'Ve el calendario de esa semana'],
        ['3', 'Hace clic en "Agregar citas"',                    'Admin',   'Se abre el formulario de slots'],
        ['4', 'Selecciona fecha, hora inicio y hora fin',        'Admin',   'Define el rango horario'],
        ['5', 'Guarda los slots',                                'Admin',   'Sistema crea los slots cada hora'],
        ['6', 'Calendario se actualiza',                         'Sistema', 'Los nuevos slots aparecen en verde'],
        ['7', 'Los clientes ya pueden ver y reservar los slots', 'Clientes', 'Disponibilidad pública'],
    ],
    anchos=[1.2, 6.3, 2.5, 6]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
#  7. COMPATIBILIDAD Y DISPOSITIVOS
# ═══════════════════════════════════════════════════════════════════
H1('7. Compatibilidad y Dispositivos')

H2('7.1 El sistema se adapta a cualquier pantalla')
body('Mara Beauty Studio está diseñado con tecnología responsive, lo que significa que se ajusta '
     'automáticamente al tamaño de pantalla del dispositivo que use.')

tabla_simple(
    ['Dispositivo', 'Tamaño de pantalla', 'Experiencia'],
    [
        ['Computador de escritorio', '1920 × 1080 px', 'Vista completa con toda la información visible'],
        ['Portátil / Laptop',        '1366 × 768 px',  'Vista completa, diseño idéntico al escritorio'],
        ['Tableta (iPad)',            '768 × 1024 px',  'Diseño adaptado, navegación táctil'],
        ['Celular iPhone',           '390 × 844 px',   'Vista vertical, contenido apilado legible'],
        ['Celular Android',          '360 × 800 px',   'Vista vertical, contenido apilado legible'],
    ],
    anchos=[4, 3.5, 8.5]
)

H2('7.2 Vista en distintos dispositivos')
img('01_responsive_Desktop_1920.png', 'Figura 10 — Vista en pantalla Desktop (1920 × 1080)', ancho=14.5)
img('01_responsive_Tablet_iPad.png',  'Figura 11 — Vista en Tableta iPad (768 × 1024)', ancho=10)
img('01_responsive_Mobile_iPhone.png','Figura 12 — Vista en iPhone 14 (390 × 844)', ancho=7)

H2('7.3 Navegadores compatibles')
tabla_simple(
    ['Navegador', 'Versión mínima', 'Estado'],
    [
        ['Google Chrome',   '100+', '✔ Totalmente compatible'],
        ['Mozilla Firefox', '100+', '✔ Totalmente compatible'],
        ['Microsoft Edge',  '100+', '✔ Totalmente compatible'],
        ['Apple Safari',    '15+',  '✔ Totalmente compatible'],
    ],
    anchos=[5, 3.5, 7.5]
)

img('17_navegador_Chrome_120.png',  'Figura 13 — Sistema en Google Chrome', ancho=14)
img('17_navegador_Firefox_121.png', 'Figura 14 — Sistema en Mozilla Firefox', ancho=14)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
#  8. PREGUNTAS FRECUENTES
# ═══════════════════════════════════════════════════════════════════
H1('8. Preguntas Frecuentes')

preguntas = [
    (
        '¿Necesito crear una cuenta para reservar una cita?',
        'No. Solo necesita su nombre y correo electrónico al momento de hacer la reserva. '
        'No existe proceso de registro previo para clientes.'
    ),
    (
        '¿Qué pasa si intento reservar un horario que otro cliente acaba de tomar?',
        'El sistema valida en tiempo real. Si el slot ya no está disponible, recibirá el mensaje '
        '"La cita ya no está disponible" y deberá elegir otro horario.'
    ),
    (
        '¿Puedo tener varias citas al mismo tiempo?',
        'Sí. Puede reservar varios slots en diferentes fechas y horarios usando el mismo correo electrónico.'
    ),
    (
        '¿Qué hago si no recibí el correo de confirmación?',
        'Revise la carpeta de spam o correo no deseado. Si el problema persiste, comuníquese '
        'directamente con el salón para verificar su reserva.'
    ),
    (
        '¿Puedo cancelar mi cita el mismo día?',
        'Sí, siempre que la cita no haya pasado ya. Use el enlace del correo de confirmación '
        'para cancelar en cualquier momento antes de la cita.'
    ),
    (
        '¿El administrador puede ver mis datos personales?',
        'Solo el nombre y correo que usted proporcionó al hacer la reserva, los cuales son necesarios '
        'para gestionar el agendamiento.'
    ),
    (
        '¿Qué ocurre si el sitio muestra error al cargar las citas?',
        'Puede ser un problema temporal de conexión. Espere unos segundos y recargue la página (F5). '
        'Si el problema persiste, el servidor del sistema puede estar fuera de servicio.'
    ),
    (
        '¿Cómo sé que mi conexión es segura?',
        'El sitio usa protocolo HTTPS (candado en la barra del navegador), lo que garantiza que '
        'la información viaja cifrada entre su dispositivo y el servidor.'
    ),
]

for i, (pregunta, respuesta) in enumerate(preguntas, 1):
    tbl = doc.add_table(rows=2, cols=1)
    tbl.style = 'Table Grid'
    cp = tbl.rows[0].cells[0]
    cr = tbl.rows[1].cells[0]
    shd(cp, 'F3E6EF'); border_cell(cp, 'C67BA3')
    shd(cr, 'FDFBFD'); border_cell(cr, 'DDDDDD')
    pp = cp.paragraphs[0]
    pp.paragraph_format.space_before = Pt(5)
    pp.paragraph_format.space_after  = Pt(5)
    pp.paragraph_format.left_indent  = Cm(0.3)
    rp = pp.add_run(f'P{i}.  {pregunta}')
    rp.bold = True; rp.font.size = Pt(10); rp.font.color.rgb = C_DARK
    pr = cr.paragraphs[0]
    pr.paragraph_format.space_before = Pt(5)
    pr.paragraph_format.space_after  = Pt(5)
    pr.paragraph_format.left_indent  = Cm(0.3)
    rr = pr.add_run(respuesta)
    rr.font.size = Pt(10); rr.font.color.rgb = C_DARK
    tbl.columns[0].width = Cm(14.4)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)

# ── PIE DE PÁGINA ────────────────────────────────────────────────
doc.add_paragraph()
fp = doc.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run(
    f'Mara Beauty Studio  ·  Manual de Usuario v1.0  ·  {datetime.now().strftime("%d/%m/%Y")}'
)
fr.font.size = Pt(8); fr.font.color.rgb = C_GRAY; fr.italic = True

doc.save(OUT)
print(f'✅  Manual de usuario generado: {OUT}')
