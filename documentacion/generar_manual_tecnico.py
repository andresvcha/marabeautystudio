#!/usr/bin/env python3
# Manual Técnico — Mara Beauty Studio
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import os

OUTPUT = '/Users/andres/Documents/Mara Beauty Studio/Manual_Tecnico_Mara_Beauty_Studio.docx'

# ── Paleta ───────────────────────────────────────────────────────────
C_DARK   = RGBColor(0x1A, 0x1A, 0x2E)
C_ROSE   = RGBColor(0xC6, 0x7B, 0xA3)
C_BLUE   = RGBColor(0x33, 0x66, 0xCC)
C_GRAY   = RGBColor(0x77, 0x77, 0x77)
C_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
C_LIGHT  = RGBColor(0xF8, 0xF4, 0xF6)

HEX_HEADER = '1A1A2E'
HEX_ROW_A  = 'FAF0F5'
HEX_ROW_B  = 'FFFFFF'
HEX_ACCENT = 'F3E6EF'

def shd(cell, hex_col):
    tc = cell._tc
    pr = tc.get_or_add_tcPr()
    s  = OxmlElement('w:shd')
    s.set(qn('w:fill'), hex_col)
    s.set(qn('w:val'), 'clear')
    pr.append(s)

def border(cell, col='DDDDDD'):
    tc = cell._tc
    pr = tc.get_or_add_tcPr()
    tb = OxmlElement('w:tcBorders')
    for side in ('top','left','bottom','right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), col)
        tb.append(b)
    pr.append(tb)

def cell_fmt(c, txt='', bold=False, size=9, color=None, bg=None, align=WD_ALIGN_PARAGRAPH.LEFT, pad=3):
    if bg:  shd(c, bg)
    border(c)
    p = c.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(pad)
    p.paragraph_format.space_after  = Pt(pad)
    if txt:
        r = p.add_run(txt)
        r.bold = bold
        r.font.size = Pt(size)
        r.font.color.rgb = color if color else C_DARK
    return p

doc = Document()

# Márgenes
for s in doc.sections:
    s.top_margin    = Cm(2.2)
    s.bottom_margin = Cm(2.2)
    s.left_margin   = Cm(2.8)
    s.right_margin  = Cm(2.8)

# Fuente base
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(10)

# ── Helpers ──────────────────────────────────────────────────────────
def H1(texto):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(texto.upper())
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = C_DARK
    # línea decorativa
    l = doc.add_paragraph()
    l.paragraph_format.space_before = Pt(0)
    l.paragraph_format.space_after  = Pt(8)
    lr = l.add_run('▬' * 60)
    lr.font.size = Pt(5)
    lr.font.color.rgb = C_ROSE

def H2(texto):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(f'◆  {texto}')
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = C_ROSE

def H3(texto):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(texto)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = C_DARK

def body(texto, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(texto)
    r.font.size = Pt(10)
    r.italic = italic
    r.font.color.rgb = C_DARK

def bullet(texto):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    r = p.add_run(texto)
    r.font.size = Pt(10)

def code_block(lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        p.paragraph_format.left_indent  = Cm(1)
        r = p.add_run(line)
        r.font.name = 'Courier New'
        r.font.size = Pt(9)
        r.font.color.rgb = C_DARK

def mono_box(lines, title=''):
    if title:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = C_ROSE
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    shd(cell, 'F0F0F0')
    border(cell, 'CCCCCC')
    cell.paragraphs[0].clear()
    for line in lines:
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_before = Pt(1)
        p2.paragraph_format.space_after  = Pt(1)
        r2 = p2.add_run(line)
        r2.font.name = 'Courier New'
        r2.font.size = Pt(8.5)
    tbl.columns[0].width = Cm(14.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def simple_table(headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell_fmt(tbl.rows[0].cells[i], h, bold=True, size=9, color=C_WHITE, bg=HEX_HEADER, align=WD_ALIGN_PARAGRAPH.CENTER, pad=4)
    for ri, row in enumerate(rows):
        bg = HEX_ROW_A if ri % 2 else HEX_ROW_B
        for ci, val in enumerate(row):
            is_mono = isinstance(val, tuple) and val[0] == 'mono'
            txt = val[1] if isinstance(val, tuple) else val
            p = cell_fmt(tbl.rows[ri+1].cells[ci], txt, size=9, bg=bg, pad=3)
            if is_mono:
                for r3 in p.runs:
                    r3.font.name = 'Courier New'
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ════════════════════════════════════════════════════════════════════
#  PORTADA
# ════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('MARA BEAUTY STUDIO')
r.font.size = Pt(28); r.bold = True; r.font.color.rgb = C_DARK

doc.add_paragraph()
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('Manual Técnico y de Usuario')
r2.font.size = Pt(17); r2.bold = True; r2.font.color.rgb = C_ROSE

doc.add_paragraph()
for line, sz, col in [
    ('Sistema de Gestión de Citas — Salón de Belleza', 11, C_GRAY),
    (f'Versión 1.0  ·  {datetime.now().strftime("%B %Y")}', 10, C_GRAY),
]:
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(line)
    r3.font.size = Pt(sz); r3.font.color.rgb = col

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#  ÍNDICE
# ════════════════════════════════════════════════════════════════════
H1('Tabla de Contenido')
toc = [
    ('1.', 'Prerrequisitos de Instalación'),
    ('2.', 'Frameworks y Estándares'),
    ('3.', 'Diagrama de Casos de Uso'),
    ('4.', 'Modelo Entidad-Relación'),
    ('5.', 'Diccionario de Datos'),
    ('6.', 'Scripts de Instalación'),
    ('7.', 'Diagrama de Componentes'),
    ('8.', 'Manual de Usuario'),
]
for num, title in toc:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(f'  {num}  {title}')
    r.font.size = Pt(11)
    r.font.color.rgb = C_DARK

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#  1. PRERREQUISITOS
# ════════════════════════════════════════════════════════════════════
H1('1. Prerrequisitos de Instalación')
body('El siguiente conjunto de herramientas y servicios debe estar instalado y configurado antes de ejecutar Mara Beauty Studio.')

H2('1.1 Software requerido')
simple_table(
    ['Herramienta', 'Versión mínima', 'Propósito', 'Descarga'],
    [
        ['Node.js',     'v18 LTS',        'Entorno de ejecución JavaScript (backend)',  'nodejs.org'],
        ['npm',         'v9+',            'Gestor de paquetes Node',                    'Incluido con Node.js'],
        ['PostgreSQL',  'v14+',           'Base de datos relacional',                   'postgresql.org'],
        ['Git',         'v2.30+',         'Control de versiones y clonado del repo',    'git-scm.com'],
        ['Navegador',   'Chrome 100+ / Firefox 100+', 'Ejecución del frontend',         'Cualquier navegador moderno'],
    ],
    col_widths=[3, 3, 5.5, 4]
)

H2('1.2 Variables de entorno requeridas')
body('Crear el archivo backend/.env con las siguientes variables:')
mono_box([
    '# ── Servidor ────────────────────────────────',
    'PORT=3000',
    'NODE_ENV=development',
    '',
    '# ── Frontend (CORS) ─────────────────────────',
    'FRONTEND_URL=http://localhost:5500',
    '',
    '# ── Base de datos ───────────────────────────',
    'DATABASE_URL="postgresql://USUARIO:CONTRASEÑA@localhost:5432/marabeauty"',
    '',
    '# ── Autenticación JWT ───────────────────────',
    'JWT_SECRET=secreto_muy_seguro_cambiar_en_produccion',
    'JWT_EXPIRES_IN=8h',
    '',
    '# ── Email (Nodemailer / Gmail) ──────────────',
    'EMAIL_HOST=smtp.gmail.com',
    'EMAIL_PORT=587',
    'EMAIL_USER=correo@gmail.com',
    'EMAIL_PASS=app_password_de_gmail',
], title='backend/.env')

H2('1.3 Puertos utilizados')
simple_table(
    ['Puerto', 'Servicio', 'Descripción'],
    [
        ['3000', 'API REST (Express.js)',   'Backend — todos los endpoints del sistema'],
        ['5432', 'PostgreSQL',              'Base de datos relacional'],
        ['5500', 'Frontend (Live Server)',  'Servidor estático del frontend en desarrollo'],
        ['4040', 'ngrok / tunnel',          'Exposición del backend en internet (opcional)'],
    ],
    col_widths=[2.5, 4.5, 9]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#  2. FRAMEWORKS Y ESTÁNDARES
# ════════════════════════════════════════════════════════════════════
H1('2. Frameworks y Estándares')

H2('2.1 Arquitectura general')
body('Mara Beauty Studio sigue una arquitectura cliente-servidor desacoplada de tres capas:')
for item in [
    'Capa de presentación: Frontend estático (HTML5 + CSS3 + JavaScript Vanilla ES6+)',
    'Capa de negocio: API REST construida con Node.js + Express.js bajo patrón MVC modular',
    'Capa de datos: PostgreSQL gestionado mediante el ORM Prisma',
]:
    bullet(item)

H2('2.2 Stack tecnológico — Backend')
simple_table(
    ['Tecnología', 'Versión', 'Rol en el sistema'],
    [
        ['Node.js',         'v18 LTS',  'Entorno de ejecución del servidor'],
        ['Express.js',      '4.18',     'Framework web — enrutamiento, middlewares, REST API'],
        ['Prisma ORM',      '5.10',     'Mapeo objeto-relacional con PostgreSQL'],
        ['JWT (jsonwebtoken)', '9.0',   'Autenticación stateless mediante tokens Bearer'],
        ['bcryptjs',        '2.4',      'Hash seguro de contraseñas (salt rounds: 10)'],
        ['Nodemailer',      '6.9',      'Envío de emails transaccionales (confirmación / cancelación)'],
        ['Swagger UI Express', '5.0',   'Documentación interactiva de la API (/api/docs)'],
        ['dotenv',          '16.4',     'Gestión de variables de entorno'],
        ['cors',            '2.8',      'Política de acceso cruzado entre dominios'],
    ],
    col_widths=[4.5, 2.5, 9]
)

H2('2.3 Stack tecnológico — Frontend')
simple_table(
    ['Tecnología', 'Versión', 'Rol en el sistema'],
    [
        ['HTML5',           'Living Standard', 'Estructura semántica de las vistas'],
        ['CSS3',            'Living Standard', 'Estilos, layout, diseño responsive'],
        ['JavaScript',      'ES2022 (ES13)',   'Lógica del cliente — fetch API, DOM, localStorage'],
        ['Google Fonts',    'Poppins',         'Tipografía corporativa del sistema'],
    ],
    col_widths=[4.5, 3.5, 8]
)

H2('2.4 Estándares y convenciones aplicadas')
simple_table(
    ['Estándar', 'Descripción'],
    [
        ['REST (Richardson Maturity Level 2)', 'Endpoints organizados por recursos, verbos HTTP semánticos (GET, POST, DELETE, PUT)'],
        ['JWT Bearer Token (RFC 7519)',         'Autenticación sin estado — token firmado HS256 con expiración de 8 horas'],
        ['OpenAPI 3.0 (Swagger)',               'Documentación automática de endpoints disponible en /api/docs'],
        ['bcrypt (NIST SP 800-132)',             'Hash de contraseñas con factor de trabajo 10 (adaptable a hardware)'],
        ['CORS (W3C)',                           'Whitelist de orígenes permitidos configurada en variables de entorno'],
        ['MVC modular',                          'Separación controller / service / routes por módulo (auth, appointments, clients, users, admin)'],
        ['Responsive Web Design',               'CSS adaptable a pantallas desde 360 px (móvil) hasta 1920 px (desktop)'],
        ['Locale es-CO / America/Bogota',        'Formateo de fechas y horas en zona horaria colombiana'],
    ],
    col_widths=[5.5, 10.5]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#  3. DIAGRAMA DE CASOS DE USO
# ════════════════════════════════════════════════════════════════════
H1('3. Diagrama de Casos de Uso')
body('El sistema contempla dos actores principales: el Cliente (usuario público sin autenticación) y el Administrador (usuario autenticado con JWT).')

doc.add_paragraph()

# Diagrama ASCII
mono_box([
    '                    ┌─────────────────────────────────────────────────────────┐',
    '                    │              SISTEMA MARA BEAUTY STUDIO                 │',
    '                    │                                                         │',
    '   ┌──────────┐     │   ┌───────────────────────────┐                        │',
    '   │          │     │   │  UC-01  Ver calendario     │                        │',
    '   │          │─────┼──▶│         de citas           │                        │',
    '   │          │     │   └───────────────────────────┘                        │',
    '   │          │     │                                                         │',
    '   │ CLIENTE  │     │   ┌───────────────────────────┐                        │',
    '   │ (público)│─────┼──▶│  UC-02  Navegar semanas   │                        │',
    '   │          │     │   └───────────────────────────┘                        │',
    '   │          │     │                                                         │',
    '   │          │     │   ┌───────────────────────────┐                        │',
    '   │          │─────┼──▶│  UC-03  Reservar cita     │                        │',
    '   └──────────┘     │   └───────────────────────────┘                        │',
    '                    │                                                         │',
    '                    │   ┌───────────────────────────┐                        │',
    '   ┌──────────┐     │   │  UC-04  Cancelar cita     │◀───────────────────┐   │',
    '   │          │─────┼──▶│         (link de email)   │                    │   │',
    '   │          │     │   └───────────────────────────┘                    │   │',
    '   │          │     │                                           «include» │   │',
    '   │          │     │   ┌─────────────────────────────────┐              │   │',
    '   │  ADMIN   │─────┼──▶│  UC-05  Iniciar sesión (Login) │──────────────┘   │',
    '   │  (auth.) │     │   └─────────────────────────────────┘                  │',
    '   │          │     │                                                         │',
    '   │          │     │   ┌───────────────────────────┐                        │',
    '   │          │─────┼──▶│  UC-06  Crear slots citas  │                        │',
    '   │          │     │   └───────────────────────────┘                        │',
    '   │          │     │                                                         │',
    '   │          │     │   ┌───────────────────────────┐                        │',
    '   │          │─────┼──▶│  UC-07  Eliminar slot     │                        │',
    '   │          │     │   └───────────────────────────┘                        │',
    '   │          │     │                                                         │',
    '   │          │     │   ┌───────────────────────────┐                        │',
    '   │          │─────┼──▶│  UC-08  Ver estadísticas  │                        │',
    '   │          │     │   └───────────────────────────┘                        │',
    '   │          │     │                                                         │',
    '   │          │     │   ┌───────────────────────────┐                        │',
    '   │          │─────┼──▶│  UC-09  Ver clientes      │                        │',
    '   │          │     │   └───────────────────────────┘                        │',
    '   │          │     │                                                         │',
    '   │          │     │   ┌───────────────────────────┐                        │',
    '   │          │─────┼──▶│  UC-10  Gestionar usuarios│                        │',
    '   └──────────┘     │   └───────────────────────────┘                        │',
    '                    └─────────────────────────────────────────────────────────┘',
    '',
    '   Leyenda:  ──▶ asociación actor-caso de uso    «include» relación de inclusión',
], title='Diagrama de Casos de Uso — representación textual')

H2('3.1 Descripción de casos de uso')
simple_table(
    ['ID', 'Caso de uso', 'Actor', 'Descripción', 'Precondición'],
    [
        ['UC-01', 'Ver calendario', 'Cliente', 'Visualiza la semana actual con slots disponibles/reservados', 'Ninguna'],
        ['UC-02', 'Navegar semanas', 'Cliente / Admin', 'Avanza o retrocede semanas en el calendario', 'Ninguna'],
        ['UC-03', 'Reservar cita', 'Cliente', 'Selecciona un slot disponible e ingresa nombre y email', 'Slot en estado AVAILABLE'],
        ['UC-04', 'Cancelar cita', 'Cliente', 'Cancela su cita usando el enlace recibido por email', 'Slot en estado RESERVED'],
        ['UC-05', 'Iniciar sesión', 'Admin', 'Autentica con usuario y contraseña; recibe token JWT', 'Usuario registrado en BD'],
        ['UC-06', 'Crear slots', 'Admin', 'Define fecha, hora inicio y fin; crea slots cada hora', 'Token JWT válido'],
        ['UC-07', 'Eliminar slot', 'Admin', 'Elimina un slot AVAILABLE del calendario', 'Token JWT válido; slot no RESERVED'],
        ['UC-08', 'Ver estadísticas', 'Admin', 'Consulta totales y conteos de citas de la semana', 'Token JWT válido'],
        ['UC-09', 'Ver clientes', 'Admin', 'Lista todos los clientes registrados con última cita', 'Token JWT válido'],
        ['UC-10', 'Gestionar usuarios', 'Admin', 'Crea nuevo admin y cambia contraseña propia', 'Token JWT válido'],
    ],
    col_widths=[1.5, 3.5, 2.5, 5.5, 3]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#  4. MODELO ENTIDAD-RELACIÓN
# ════════════════════════════════════════════════════════════════════
H1('4. Modelo Entidad-Relación')
body('La base de datos marabeauty contiene tres entidades principales definidas en el schema Prisma.')

mono_box([
    '  ┌──────────────────────┐         ┌───────────────────────────────────────┐',
    '  │        USER          │         │            APPOINTMENT                │',
    '  ├──────────────────────┤         ├───────────────────────────────────────┤',
    '  │ PK  id         INT   │         │ PK  id          INT                   │',
    '  │     username   STRING│         │     dateTime    DATETIME  (UNIQUE)    │',
    '  │     password   STRING│         │     state       ENUM                  │',
    '  │     createdAt  DATE  │         │               (AVAILABLE | RESERVED)  │',
    '  └──────────────────────┘         │ FK  clientId    INT?  ────────────┐   │',
    '                                   │     createdAt   DATETIME          │   │',
    '                                   │     updatedAt   DATETIME          │   │',
    '                                   └───────────────────────────────────┼───┘',
    '                                                                        │',
    '                                                                        │  0..N',
    '                                                              1 ────────┘',
    '                                   ┌──────────────────────┐',
    '                                   │        CLIENT        │',
    '                                   ├──────────────────────┤',
    '                                   │ PK  id        INT    │',
    '                                   │     name      STRING │',
    '                                   │     email     STRING │ (UNIQUE)',
    '                                   │     createdAt DATE   │',
    '                                   └──────────────────────┘',
    '',
    '  Relaciones:',
    '    CLIENT ──┤ tiene ├──< APPOINTMENT   (uno a muchos; clientId nullable)',
    '    USER  — sin relaciones directas (entidad independiente de autenticación)',
    '',
    '  Reglas de negocio:',
    '    • Un APPOINTMENT.dateTime es único — no pueden existir dos citas en el mismo horario',
    '    • clientId es NULL cuando state = AVAILABLE; no-nulo cuando state = RESERVED',
    '    • Un CLIENT puede tener múltiples APPOINTMENTs a lo largo del tiempo',
], title='Diagrama Entidad-Relación')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#  5. DICCIONARIO DE DATOS
# ════════════════════════════════════════════════════════════════════
H1('5. Diccionario de Datos')

H2('5.1 Tabla: User')
body('Almacena los administradores del sistema. Gestionada solo por endpoints protegidos con JWT.')
simple_table(
    ['Campo', 'Tipo SQL', 'Tipo Prisma', 'Restricción', 'Descripción'],
    [
        [('mono','id'),        ('mono','SERIAL'),     ('mono','Int'),      'PK, AUTO_INCREMENT',   'Identificador único del usuario'],
        [('mono','username'),  ('mono','VARCHAR'),    ('mono','String'),   'UNIQUE, NOT NULL',     'Nombre de usuario para el login'],
        [('mono','password'),  ('mono','TEXT'),       ('mono','String'),   'NOT NULL',             'Hash bcrypt de la contraseña (salt=10)'],
        [('mono','createdAt'), ('mono','TIMESTAMPTZ'),('mono','DateTime'), 'DEFAULT now()',        'Fecha y hora de creación del registro'],
    ],
    col_widths=[2.8, 3, 2.8, 3.5, 5]
)

H2('5.2 Tabla: Client')
body('Registra los clientes que realizan reservas. Se crea automáticamente al reservar una cita.')
simple_table(
    ['Campo', 'Tipo SQL', 'Tipo Prisma', 'Restricción', 'Descripción'],
    [
        [('mono','id'),        ('mono','SERIAL'),     ('mono','Int'),      'PK, AUTO_INCREMENT',   'Identificador único del cliente'],
        [('mono','name'),      ('mono','VARCHAR'),    ('mono','String'),   'NOT NULL',             'Nombre completo del cliente'],
        [('mono','email'),     ('mono','VARCHAR'),    ('mono','String'),   'UNIQUE, NOT NULL',     'Email del cliente — clave de búsqueda'],
        [('mono','createdAt'), ('mono','TIMESTAMPTZ'),('mono','DateTime'), 'DEFAULT now()',        'Fecha de primer registro'],
    ],
    col_widths=[2.8, 3, 2.8, 3.5, 5]
)

H2('5.3 Tabla: Appointment')
body('Entidad central del sistema. Representa cada slot de cita con su estado y cliente asignado.')
simple_table(
    ['Campo', 'Tipo SQL', 'Tipo Prisma', 'Restricción', 'Descripción'],
    [
        [('mono','id'),        ('mono','SERIAL'),     ('mono','Int'),      'PK, AUTO_INCREMENT',   'Identificador único de la cita'],
        [('mono','dateTime'),  ('mono','TIMESTAMPTZ'),('mono','DateTime'), 'UNIQUE, NOT NULL',     'Fecha y hora exacta del slot (sin solapamiento)'],
        [('mono','state'),     ('mono','ENUM'),       ('mono','Enum'),     'DEFAULT AVAILABLE',    'Estado: AVAILABLE (libre) o RESERVED (tomada)'],
        [('mono','clientId'),  ('mono','INTEGER'),    ('mono','Int?'),     'FK → Client.id, NULL', 'Referencia al cliente; NULL si está disponible'],
        [('mono','createdAt'), ('mono','TIMESTAMPTZ'),('mono','DateTime'), 'DEFAULT now()',        'Fecha de creación del slot'],
        [('mono','updatedAt'), ('mono','TIMESTAMPTZ'),('mono','DateTime'), 'AUTO-UPDATE',          'Última modificación (Prisma lo gestiona)'],
    ],
    col_widths=[2.8, 3, 2.8, 3.5, 5]
)

H2('5.4 Enum: AppointmentState')
simple_table(
    ['Valor', 'Significado', 'Transición permitida'],
    [
        ['AVAILABLE', 'Slot libre — visible en el calendario para reserva',       'AVAILABLE → RESERVED (al reservar)'],
        ['RESERVED',  'Slot ocupado por un cliente — no eliminable por el admin', 'RESERVED → AVAILABLE (al cancelar)'],
    ],
    col_widths=[3, 6.5, 6.5]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#  6. SCRIPTS DE INSTALACIÓN
# ════════════════════════════════════════════════════════════════════
H1('6. Scripts de Instalación')
body('Siga los pasos en orden. Todos los comandos se ejecutan desde la terminal en la raíz del proyecto.')

H2('6.1 Clonar el repositorio')
mono_box([
    '# Clonar desde GitHub',
    'git clone https://github.com/tu-usuario/mara-beauty-studio.git',
    'cd mara-beauty-studio',
], title='Terminal')

H2('6.2 Configurar el Backend')
mono_box([
    '# 1. Entrar a la carpeta del backend',
    'cd backend',
    '',
    '# 2. Instalar dependencias',
    'npm install',
    '',
    '# 3. Crear el archivo de variables de entorno',
    'cp .env.example .env',
    '# Editar .env con los valores reales (DB, JWT_SECRET, email)',
    '',
    '# 4. Crear la base de datos en PostgreSQL',
    'psql -U postgres -c "CREATE DATABASE marabeauty;"',
    '',
    '# 5. Ejecutar migraciones (crea las tablas)',
    'npx prisma migrate dev --name init',
    '',
    '# 6. Cargar datos iniciales (usuario admin)',
    'node prisma/seed.js',
    '',
    '# 7. Iniciar el servidor',
    'node server.js',
    '# → Servidor corriendo en http://localhost:3000',
    '# → Documentación API: http://localhost:3000/api/docs',
], title='Instalación Backend')

H2('6.3 Servir el Frontend')
mono_box([
    '# Opción A — Extensión Live Server (VS Code)',
    '# Abrir la carpeta frontend/ en VS Code y clic en "Go Live"',
    '# → Corre en http://localhost:5500',
    '',
    '# Opción B — Servidor HTTP simple (Python)',
    'cd frontend',
    'python3 -m http.server 5500',
    '',
    '# Opción C — npx serve',
    'npx serve frontend -p 5500',
], title='Servir Frontend')

H2('6.4 Verificar la instalación')
mono_box([
    '# Health check del backend',
    'curl http://localhost:3000/api/health',
    '# → {"status":"OK","message":"Mara Beauty Studio API activa"}',
    '',
    '# Probar login',
    'curl -X POST http://localhost:3000/api/auth/login \\',
    '     -H "Content-Type: application/json" \\',
    '     -d \'{"username":"admin","password":"admin123"}\'',
    '# → {"token":"eyJ...","user":{"id":1,"username":"admin"}}',
], title='Verificación')

H2('6.5 Comandos útiles de Prisma')
simple_table(
    ['Comando', 'Descripción'],
    [
        [('mono','npx prisma migrate dev'),    'Aplica migraciones pendientes y regenera el cliente Prisma'],
        [('mono','npx prisma db push'),        'Sincroniza el schema sin crear archivos de migración (desarrollo rápido)'],
        [('mono','npx prisma studio'),         'Abre interfaz web para explorar y editar la base de datos'],
        [('mono','npx prisma generate'),       'Regenera el cliente Prisma sin migrar (tras cambios en schema.prisma)'],
        [('mono','node prisma/seed.js'),       'Carga datos iniciales: usuario admin con contraseña admin123'],
    ],
    col_widths=[6, 10]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#  7. DIAGRAMA DE COMPONENTES
# ════════════════════════════════════════════════════════════════════
H1('7. Diagrama de Componentes')

mono_box([
    '┌─────────────────────────────────────────────────────────────────────────────┐',
    '│                        CLIENTE (Navegador Web)                              │',
    '│  ┌──────────────┐  ┌──────────────┐  ┌───────────────┐  ┌──────────────┐  │',
    '│  │  index.html  │  │  admin.html  │  │ dashboard.html│  │simulador.html│  │',
    '│  │  (Calendario │  │  (Login      │  │  (Panel Admin │  │  (Preview    │  │',
    '│  │   cliente)   │  │   Admin)     │  │   Dashboard)  │  │   Emails)    │  │',
    '│  └──────┬───────┘  └──────┬───────┘  └───────┬───────┘  └──────────────┘  │',
    '│         │                 │                  │                             │',
    '│  ┌──────▼─────────────────▼──────────────────▼──────────────────────────┐  │',
    '│  │                    JavaScript Modules (ES6+)                         │  │',
    '│  │  ┌────────┐  ┌────────┐  ┌──────────┐  ┌─────────┐  ┌───────────┐  │  │',
    '│  │  │ api.js │  │auth.js │  │calendar.js│  │client.js│  │dashboard.js│  │  │',
    '│  │  │(fetch/ │  │(login/ │  │(render/  │  │(reservar│  │(crear/    │  │  │',
    '│  │  │ token) │  │ logout)│  │ navegac.)│  │ cancelar│  │ eliminar) │  │  │',
    '│  │  └────────┘  └────────┘  └──────────┘  └─────────┘  └───────────┘  │  │',
    '│  │  ┌────────┐  ┌────────┐                                             │  │',
    '│  │  │  ui.js │  │utils.js│   localStorage: mara_token (JWT)            │  │',
    '│  │  └────────┘  └────────┘                                             │  │',
    '│  └────────────────────────────────────┬─────────────────────────────────┘  │',
    '└───────────────────────────────────────┼─────────────────────────────────────┘',
    '                                        │ HTTPS / REST API (JSON)',
    '                                        ▼',
    '┌─────────────────────────────────────────────────────────────────────────────┐',
    '│                      SERVIDOR (Node.js / Express)                           │',
    '│                                                                             │',
    '│  ┌──────────────┐    ┌──────────────────────────────────────────────────┐  │',
    '│  │  server.js   │    │                    src/app.js                    │  │',
    '│  │  (Entry Point│───▶│  Express App — CORS · JSON · Swagger · Errores  │  │',
    '│  │   + dotenv)  │    └────────────────────┬─────────────────────────────┘  │',
    '│  └──────────────┘                         │                               │',
    '│  ┌────────────────────────────────────────▼──────────────────────────────┐ │',
    '│  │                         MÓDULOS (Routes → Controller → Service)       │ │',
    '│  │  ┌─────────┐  ┌──────────────┐  ┌─────────┐  ┌────────┐  ┌────────┐ │ │',
    '│  │  │  /auth  │  │/appointments │  │ /admin  │  │/clients│  │ /users │ │ │',
    '│  │  │ (login) │  │(get/post/del │  │(stats / │  │ (list) │  │(profile│ │ │',
    '│  │  │         │  │ /book/cancel)│  │ clients)│  │        │  │ /pass) │ │ │',
    '│  │  └────┬────┘  └──────┬───────┘  └────┬────┘  └───┬────┘  └────┬───┘ │ │',
    '│  └───────┼──────────────┼───────────────┼────────────┼────────────┼─────┘ │',
    '│          │  JWT         │               │            │            │        │',
    '│  ┌───────▼──────────────▼───────────────▼────────────▼────────────▼─────┐  │',
    '│  │              Middleware: auth.middleware · error.middleware           │  │',
    '│  └────────────────────────────┬──────────────────────────────────────────┘  │',
    '│                               │                                            │',
    '│  ┌────────────────────────────▼──────────┐  ┌─────────────────────────┐   │',
    '│  │           Prisma Client (ORM)          │  │  Nodemailer (SMTP)      │   │',
    '│  │   prisma.appointment / client / user   │  │  Confirmación reserva   │   │',
    '│  └────────────────────────────┬───────────┘  │  Cancelación cita       │   │',
    '│                               │              └─────────────┬───────────┘   │',
    '└───────────────────────────────┼────────────────────────────┼───────────────┘',
    '                                │                            │ SMTP TLS',
    '                                ▼                            ▼',
    '┌───────────────────────────────────┐       ┌───────────────────────────────┐',
    '│   PostgreSQL 14+  (marabeauty)    │       │   Servidor SMTP (Gmail)       │',
    '│  • User                           │       │   smtp.gmail.com:587          │',
    '│  • Client                         │       │   Envío de notificaciones     │',
    '│  • Appointment                    │       └───────────────────────────────┘',
    '└───────────────────────────────────┘',
], title='Diagrama de Componentes')

H2('7.1 Descripción de componentes')
simple_table(
    ['Componente', 'Tecnología', 'Responsabilidad'],
    [
        ['index.html',       'HTML5 + CSS3',      'Vista del cliente — calendario semanal, formulario de reserva y cancelación'],
        ['admin.html',       'HTML5 + CSS3',      'Formulario de autenticación del administrador'],
        ['dashboard.html',   'HTML5 + CSS3',      'Panel de gestión: visualización y administración de citas'],
        ['api.js',           'JavaScript ES6',    'Capa de comunicación HTTP — fetch + headers JWT + manejo de errores'],
        ['auth.js',          'JavaScript ES6',    'Login, logout, redirección y persistencia del token en localStorage'],
        ['calendar.js',      'JavaScript ES6',    'Renderizado dinámico del calendario por semana'],
        ['client.js',        'JavaScript ES6',    'Reservas y cancelaciones desde la vista del cliente'],
        ['dashboard.js',     'JavaScript ES6',    'CRUD de slots desde la vista del administrador'],
        ['server.js',        'Node.js',           'Punto de entrada: carga dotenv e inicia Express en el puerto configurado'],
        ['app.js',           'Express.js',        'Configuración central: CORS, JSON, Swagger, rutas y manejo de errores'],
        ['auth.routes',      'Express Router',    'POST /api/auth/login — emite token JWT'],
        ['appointments.routes','Express Router',  'CRUD de citas: GET (semana), POST (crear slots), DELETE, POST /:id/book, POST /:id/cancel'],
        ['admin.routes',     'Express Router',    'GET /api/admin/stats y /api/admin/clients — solo admin autenticado'],
        ['users.routes',     'Express Router',    'GET profile, POST crear admin, PUT cambiar contraseña'],
        ['auth.middleware',  'Express Middleware', 'verifyToken — valida JWT Bearer en rutas protegidas'],
        ['notifications.service','Nodemailer',    'Envío asíncrono de emails de confirmación y cancelación de cita'],
        ['Prisma Client',    'Prisma ORM',        'Abstracción de BD — consultas tipadas a PostgreSQL sin SQL directo'],
        ['PostgreSQL',       'Base de datos',     'Almacén persistente de User, Client y Appointment'],
    ],
    col_widths=[3.8, 3, 9.2]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
#  8. MANUAL DE USUARIO
# ════════════════════════════════════════════════════════════════════
H1('8. Manual de Usuario')
body('Esta sección describe el funcionamiento del sistema desde la perspectiva de cada actor.')

H2('8.1 Cliente — Reservar una cita')
simple_table(
    ['Paso', 'Acción', 'Resultado esperado'],
    [
        ['1', 'Ingresar a la URL del sistema en el navegador',                   'Se carga el calendario con la semana actual'],
        ['2', 'Usar los botones < > para navegar entre semanas',                 'El calendario muestra los slots de la semana seleccionada'],
        ['3', 'Hacer clic en un slot verde (DISPONIBLE)',                        'Se resalta el slot y se activa el botón "Reservar cita"'],
        ['4', 'Hacer clic en "Reservar cita" — ingresar nombre y email',         'Se abre el formulario de reserva'],
        ['5', 'Completar los campos y hacer clic en "Confirmar"',                'La cita queda en estado RESERVADA y se envía email de confirmación'],
    ],
    col_widths=[1.2, 7, 7.8]
)

H2('8.2 Cliente — Cancelar una cita')
simple_table(
    ['Paso', 'Acción', 'Resultado esperado'],
    [
        ['1', 'Abrir el email de confirmación recibido',                          'Email con enlace de cancelación'],
        ['2', 'Hacer clic en el enlace "Cancelar Cita"',                          'Se abre el sitio con el modal de cancelación activo'],
        ['3', 'Confirmar la cancelación en el modal',                             'El slot vuelve a estado DISPONIBLE y llega email de cancelación'],
    ],
    col_widths=[1.2, 7, 7.8]
)

H2('8.3 Administrador — Gestionar el sistema')
simple_table(
    ['Tarea', 'Pasos', 'Ruta'],
    [
        ['Iniciar sesión',      'Ir a /admin.html → ingresar usuario y contraseña → clic en "Iniciar sesión"',    '/admin.html'],
        ['Crear slots de citas', 'Dashboard → clic "Agregar citas" → seleccionar fecha, hora inicio y fin → Guardar','/dashboard.html'],
        ['Eliminar un slot',    'Dashboard → clic en slot AVAILABLE → clic "Eliminar" → Confirmar',               '/dashboard.html'],
        ['Ver estadísticas',    'Se muestran automáticamente en el dashboard: total, reservadas, disponibles',    '/dashboard.html'],
        ['Cerrar sesión',       'Clic en "Cerrar sesión" — el token se borra de localStorage',                   '/dashboard.html'],
    ],
    col_widths=[3.5, 9.5, 3]
)

H2('8.4 Preguntas frecuentes')
for q, a in [
    ('¿Qué pasa si dos personas intentan reservar el mismo slot al mismo tiempo?',
     'El sistema valida el estado en la base de datos antes de confirmar. El segundo intento recibirá un error 409 "La cita ya no está disponible".'),
    ('¿El email de confirmación es obligatorio para reservar?',
     'No. Si el servicio de email no está configurado (EMAIL_USER vacío), la cita se reserva igualmente pero sin enviar notificación.'),
    ('¿Puedo usar el sistema desde el celular?',
     'Sí. El diseño es responsive y se adapta a pantallas desde 360 px (móvil) hasta 1920 px (desktop).'),
    ('¿Cómo cambio la contraseña del administrador?',
     'Desde el dashboard, mediante el endpoint PUT /api/users/password con la contraseña actual y la nueva.'),
]:
    H3(f'• {q}')
    body(f'  {a}')

# ── Pie ────────────────────────────────────────────────────────────
doc.add_paragraph()
fp = doc.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run(f'Mara Beauty Studio  ·  Manual Técnico v1.0  ·  {datetime.now().strftime("%d/%m/%Y")}')
fr.font.size  = Pt(8)
fr.font.color.rgb = C_GRAY
fr.italic = True

doc.save(OUTPUT)
print(f'✅  Manual generado: {OUTPUT}')
