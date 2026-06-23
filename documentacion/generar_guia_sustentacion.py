from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Márgenes
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

NARANJA = RGBColor(0xF0, 0x8C, 0x00)
OSCURO  = RGBColor(0x1A, 0x1A, 0x1A)
GRIS    = RGBColor(0x55, 0x55, 0x55)
VERDE   = RGBColor(0x1A, 0x6B, 0x3C)

def estilo_normal(doc):
    estilo = doc.styles['Normal']
    estilo.font.name = 'Calibri'
    estilo.font.size = Pt(11)
    estilo.font.color.rgb = OSCURO

def titulo_documento(doc, texto):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(texto)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = NARANJA
    run.font.name = 'Calibri'

def subtitulo(doc, texto):
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = NARANJA
    run.font.name = 'Calibri'
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(4)

def seccion(doc, texto):
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = OSCURO
    run.font.name = 'Calibri'
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)

def etiqueta(doc, label, contenido):
    p = doc.add_paragraph()
    r1 = p.add_run(label + ": ")
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = VERDE
    r1.font.name = 'Calibri'
    r2 = p.add_run(contenido)
    r2.font.size = Pt(11)
    r2.font.color.rgb = OSCURO
    r2.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(2)

def parrafo(doc, texto, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.font.size = Pt(11)
    run.font.color.rgb = OSCURO
    run.font.name = 'Calibri'
    run.italic = italic
    p.paragraph_format.space_after = Pt(4)

def nota(doc, texto):
    p = doc.add_paragraph()
    run = p.add_run("ENFASIS: " + texto)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0xB0, 0x00, 0x00)
    run.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(4)

def separador(doc):
    doc.add_paragraph()

estilo_normal(doc)

# PORTADA
titulo_documento(doc, "GUÍA DE SUSTENTACIÓN")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Sistema de información web para la gestión de citas\nMara Beauty Studio")
r.font.size = Pt(13)
r.font.color.rgb = GRIS
r.font.name = 'Calibri'
separador(doc)
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Andrés Isaac Velandia Chacón  |  Ficha 3070229")
r2.font.size = Pt(11)
r2.font.color.rgb = GRIS
r2.font.name = 'Calibri'

doc.add_page_break()

# INTRODUCCIÓN
subtitulo(doc, "CÓMO USAR ESTA GUÍA")
parrafo(doc, "Este documento te indica qué decir en cada diapositiva y qué mostrar en el código. Léelo en voz alta varias veces antes de la sustentación para que fluya natural. No memorices palabra por palabra, interioriza la idea y exprésala con tus propias palabras.")
parrafo(doc, "La sustentación tiene tres momentos:")
parrafo(doc, "  1. Las diapositivas (parte documental): explicas el proyecto, el problema y la solución.")
parrafo(doc, "  2. El código en Visual Studio Code (parte técnica): muestras cómo está construido.")
parrafo(doc, "  3. La demo en vivo: abres la aplicación y registras una cita frente a los instructores.")

doc.add_page_break()

# PARTE 1: DIAPOSITIVAS
subtitulo(doc, "PARTE 1 — LAS DIAPOSITIVAS")
parrafo(doc, "Sigue el orden de tu presentación. Para cada diapositiva encontrarás qué decir y qué enfatizar.")

# SLIDE 1
seccion(doc, "Diapositiva 1 — Portada")
etiqueta(doc, "Qué decir", "\"Buenos días / tardes. Mi nombre es Andrés Isaac Velandia Chacón. Voy a presentar el proyecto Sistema de información web para la gestión de citas en Mara Beauty Studio, desarrollado en la ficha 3070229 del Centro de Comercio y Servicios, bajo la orientación de los instructores Jairo España y Juan José Murillo.\"")

# SLIDE 2
seccion(doc, "Diapositiva 2 — Agenda de trabajo")
etiqueta(doc, "Qué decir", "\"Esta es la estructura de nuestra presentación. Vamos a recorrer el título, los objetivos, el problema que identificamos, la justificación del proyecto, el alcance de lo que construimos, el impacto que genera, el trabajo futuro y los avances entregados.\"")

# SLIDE 3
seccion(doc, "Diapositiva 3 — Título del proyecto")
etiqueta(doc, "Qué decir", "\"El proyecto se llama: Sistema de información web para la gestión de citas en Mara Beauty Studio. Es una aplicación web que permite a los clientes del salón agendar, consultar y cancelar sus citas desde cualquier dispositivo con internet.\"")
nota(doc, "Enfatiza que es una aplicación web completa, no solo una página estática.")

# SLIDE 4
seccion(doc, "Diapositiva 4 — Objetivo general")
etiqueta(doc, "Qué decir", "\"El objetivo general fue construir un sistema de información web para sistematizar la gestión de citas del salón uñas Mara Beauty Studio. Es decir, reemplazar el proceso manual y desordenado por una solución digital que funcione de forma organizada y automática.\"")
nota(doc, "La palabra clave es SISTEMATIZAR: pasar del desorden manual a un proceso digital controlado.")

# SLIDE 5
seccion(doc, "Diapositiva 5 — Objetivos específicos")
etiqueta(doc, "OE1", "\"Implementar una interfaz web donde los clientes puedan ver los horarios disponibles, reservar una cita y cancelarla si lo necesitan, todo desde su teléfono o computador.\"")
etiqueta(doc, "OE2", "\"Construir un panel de administración donde la dueña del salón pueda ver todas las citas, qué horarios están ocupados, cuáles están libres, y gestionar esa disponibilidad en tiempo real.\"")
etiqueta(doc, "OE3", "\"Automatizar el envío de correos electrónicos de recordatorio, para que el cliente reciba una notificación antes de su cita y no la olvide, reduciendo las ausencias.\"")
nota(doc, "Conecta cada objetivo con algo concreto que puedas mostrar en la demo.")

# SLIDE 6
seccion(doc, "Diapositiva 6 — Planteamiento del problema (introducción)")
etiqueta(doc, "Qué decir", "\"Mara Beauty Studio es un salón de uñas ubicado en Bogotá, Colombia. A pesar de prestar un buen servicio, tenía un problema serio en la forma en que manejaba las citas. Todo se hacía a mano, con una agenda física. Eso generaba varias dificultades que vamos a ver a continuación.\"")

# SLIDE 7
seccion(doc, "Diapositiva 7 — Problema: Choques de horario")
etiqueta(doc, "Qué decir", "\"El primer problema eran los choques de horario. Como todo se anotaba a mano, era fácil cometer el error de asignarle el mismo horario a dos clientas distintas. Cuando las dos llegaban al salón al mismo tiempo, se generaba un conflicto que afectaba la experiencia de ambas y la imagen del salón.\"")
nota(doc, "Puedes decir: 'Imagínense dos personas llegando a la vez por el mismo turno — eso es exactamente lo que pasaba.'")

# SLIDE 8
seccion(doc, "Diapositiva 8 — Problema: Ausentismo")
etiqueta(doc, "Qué decir", "\"El segundo problema era el ausentismo. Sin ningún recordatorio, muchas clientas sencillamente olvidaban su cita. Ese espacio quedaba vacío y el salón perdía tiempo productivo que ya no podía recuperar.\"")
nota(doc, "Este problema lo resuelve directamente el módulo de notificaciones por correo — conéctalo.")

# SLIDE 9
seccion(doc, "Diapositiva 9 — Problema: Dependencia del horario comercial")
etiqueta(doc, "Qué decir", "\"El tercer problema era que las clientas solo podían reservar cuando alguien en el salón pudiera contestar el teléfono o responder el WhatsApp. Si llamaban a las 10 de la noche o un domingo, simplemente no podían agendar. Con el sistema web, ese límite desaparece: se puede reservar a cualquier hora del día.\"")

# SLIDE 10-12
seccion(doc, "Diapositivas 10, 11 y 12 — Justificación")
etiqueta(doc, "Qué decir (slide 10)", "\"La primera razón por la que construimos este sistema fue eliminar el proceso manual. La agenda física generaba errores, imprecisiones y pérdida de tiempo en un proceso que hoy exige soluciones digitales.\"")
etiqueta(doc, "Qué decir (slide 11)", "\"La segunda razón fue mejorar la experiencia del cliente. Con el sistema, la clienta puede entrar a la página, ver los horarios disponibles y reservar en el momento que quiera, sin depender de nadie.\"")
etiqueta(doc, "Qué decir (slide 12)", "\"La tercera razón fue reducir el ausentismo. El sistema envía automáticamente un correo de recordatorio cuando la cita se acerca, lo que disminuye las ausencias y protege la productividad del salón.\"")

# SLIDE 13
seccion(doc, "Diapositiva 13 — Alcances")
etiqueta(doc, "Qué decir", "\"El sistema hace cinco cosas: muestra los horarios disponibles en tiempo real, permite reservar una cita, permite cancelarla, tiene un panel para que el administrador gestione la disponibilidad, y envía notificaciones por correo. Lo que el sistema no hace: no asigna empleados, no guarda historial detallado de clientes, no permite reagendar directamente, y no procesa pagos en línea. Estas son funciones que quedaron definidas para una versión futura.\"")
nota(doc, "Al mencionar lo que NO hace, explica siempre que eso queda en el trabajo futuro — no es una limitación, es una decisión de alcance.")

# SLIDE 14
seccion(doc, "Diapositiva 14 — Impacto")
etiqueta(doc, "SOCIAL", "\"El sistema mejora la experiencia de las clientas. Ahora pueden gestionar sus citas desde el celular, sin llamar, sin esperar. Eso genera satisfacción y fidelización.\"")
etiqueta(doc, "AMBIENTAL", "\"Se elimina el uso de papel. Ya no se necesita la agenda física ni formularios impresos. Eso reduce el consumo de recursos físicos.\"")
etiqueta(doc, "ECONÓMICO", "\"Al reducir las ausencias y organizar mejor los horarios, el salón aprovecha mejor su tiempo disponible. Menos tiempos muertos significa más ingresos potenciales.\"")

# SLIDE 15
seccion(doc, "Diapositiva 15 — Trabajo futuro")
etiqueta(doc, "Qué decir", "\"Estas son las funcionalidades que quedaron definidas para una versión 2.0 del sistema. La asignación de empleados por servicio, los perfiles de clientes con historial de visitas, un catálogo de servicios con precios y duración, y la opción de que las clientas dejen reseñas. Todo esto está identificado y documentado para el siguiente ciclo de desarrollo.\"")
nota(doc, "Esto demuestra visión a futuro — los instructores lo valoran positivamente.")

# SLIDE 16
seccion(doc, "Diapositiva 16 — Avances del proyecto")
etiqueta(doc, "Qué decir", "\"Estos son los entregables técnicos del proyecto: el diagrama de casos de uso que muestra cómo interactúan los usuarios con el sistema, el modelo entidad relación de la base de datos, el diccionario de datos, el diagrama de componentes, el diagrama de clases, la aplicación web completa con frontend, backend y base de datos, y la documentación de la API.\"")
nota(doc, "Si los instructores preguntan por alguno de estos documentos, ten listos los archivos para mostrar.")

doc.add_page_break()

# PARTE 2: TÉCNICA
subtitulo(doc, "PARTE 2 — LA PARTE TÉCNICA EN VISUAL STUDIO CODE")
parrafo(doc, "Aquí no hay diapositivas. Abres Visual Studio Code y vas explicando la estructura mientras la muestras. Sigue este orden exacto.")

separador(doc)

# FRONTEND
seccion(doc, "2.1 — FRONTEND (lo que ve el usuario)")
parrafo(doc, "Abre la carpeta 'frontend' en VS Code y muestra su contenido.")
etiqueta(doc, "Qué decir", "\"El frontend es la parte visual de la aplicación, lo que ve y usa el cliente desde su navegador. Está construido con tecnologías web básicas: HTML, CSS y JavaScript puro. No usa ningún framework de interfaz como React, Angular o Vue.\"")

seccion(doc, "¿Por qué no se usó un framework en el frontend?")
parrafo(doc, "Si los instructores preguntan esto, responde así:")
etiqueta(doc, "Respuesta", "\"Tomamos la decisión de no usar un framework de frontend porque el alcance del proyecto es concreto y bien definido: una interfaz de calendario, un formulario de reserva y un panel de administración. Un framework como React o Angular está diseñado para aplicaciones muy grandes con muchas partes que cambian constantemente. Usarlo aquí habría agregado complejidad innecesaria. Con HTML, CSS y JavaScript puro logramos exactamente lo que necesitábamos: una aplicación funcional, liviana, que carga rápido y que cualquier navegador puede ejecutar sin configuración adicional. Además, el código es más fácil de leer y de explicar, porque no hay capas intermedias que oculten lo que está pasando.\"")
nota(doc, "PUNTO A FAVOR: No usar framework no es una limitación — es una decisión técnica. Demuestra que entiendes cuándo una herramienta es necesaria y cuándo no.")
parrafo(doc, "Muestra los archivos:")
parrafo(doc, "  • index.html → \"Esta es la página principal que ve el cliente cuando entra al sistema.\"")
parrafo(doc, "  • dashboard.html → \"Esta es la vista que ve el administrador después de iniciar sesión.\"")
parrafo(doc, "  • admin.html → \"Este es el panel de administración donde se gestiona la disponibilidad.\"")
parrafo(doc, "  • css/ → \"Aquí están los estilos visuales, los colores, las fuentes, los diseños de cada página.\"")
parrafo(doc, "  • js/api.js → \"Este es el archivo más importante del frontend. Aquí está la lógica que conecta el frontend con el backend.\"")
parrafo(doc, "La carpeta js/ tiene 7 archivos. Cada uno tiene una función específica — no todo está mezclado en un solo lugar. Eso hace el código más organizado y fácil de entender.")

etiqueta(doc, "api.js", "\"Este es el archivo base de comunicación. Contiene la función que se encarga de enviar y recibir información del servidor. Cada vez que el sistema necesita consultar algo o guardar algo, pasa por aquí. También es el que agrega el token de seguridad en cada petición, para que el servidor sepa quién está haciendo la solicitud.\"")

etiqueta(doc, "auth.js", "\"Este archivo maneja el inicio de sesión. Cuando el administrador escribe su usuario y contraseña y hace clic en entrar, este archivo toma esos datos, los envía al servidor y, si son correctos, guarda el token en el navegador y redirige al panel de administración. Si las credenciales son incorrectas, muestra un mensaje de error.\"")

etiqueta(doc, "calendar.js", "\"Este es el archivo más importante del frontend. Construye y dibuja el calendario semanal que ven tanto los clientes como el administrador. Muestra los horarios disponibles y ocupados, permite navegar entre semanas y actualiza la vista automáticamente cuando hay cambios. Comparte su información con los otros archivos para que todos trabajen con los mismos datos.\"")

etiqueta(doc, "client.js", "\"Este archivo controla todo lo que puede hacer el cliente: seleccionar un horario disponible, abrir el formulario de reserva, confirmar la cita y también cancelarla. Tiene además una función especial: si el cliente llega a la página desde el enlace del correo de recordatorio con intención de cancelar, este archivo detecta eso automáticamente y abre el formulario de cancelación de forma directa.\"")

etiqueta(doc, "dashboard.js", "\"Este archivo controla el panel del administrador. Permite crear nuevos horarios disponibles, eliminar citas existentes y cerrar sesión. Solo funciona si hay un token válido — si alguien intenta entrar sin haber iniciado sesión, lo redirige automáticamente al login.\"")

etiqueta(doc, "ui.js", "\"Este archivo maneja todo lo visual que aparece y desaparece en pantalla: los modales, que son las ventanas emergentes de confirmación, y las alertas de éxito o error. No contiene lógica del negocio — su único trabajo es mostrar o esconder elementos visuales en el momento correcto.\"")

etiqueta(doc, "utils.js", "\"Este archivo contiene funciones de apoyo para trabajar con fechas. Por ejemplo, calcula a qué número de semana del año corresponde una fecha, o convierte un número de semana en una fecha real. Son cálculos que se usan en el calendario y que, al estar separados aquí, se pueden reutilizar sin repetir código.\"")

nota(doc, "PUNTO CLAVE: El frontend no tiene lógica del negocio — solo muestra datos y envía peticiones. Toda la lógica real está en el backend. Esta separación es una buena práctica de desarrollo.")

separador(doc)

# BACKEND
seccion(doc, "2.2 — BACKEND (el motor del sistema)")
parrafo(doc, "Abre la carpeta 'backend/src' en VS Code.")
etiqueta(doc, "Framework utilizado", "Node.js + Express.js versión 4.18.2")
etiqueta(doc, "Qué decir", "\"El backend está construido sobre dos tecnologías principales. La primera es Node.js, que es el entorno de ejecución: nos permite correr JavaScript fuera del navegador, directamente en el servidor. La segunda es Express.js, que es el framework — es decir, la estructura que organiza cómo el servidor recibe, procesa y responde las peticiones. Express es el framework más usado del mundo para construir servidores con Node.js.\"")

seccion(doc, "¿Por qué Express.js?")
parrafo(doc, "Si los instructores preguntan esto, responde así:")
etiqueta(doc, "Respuesta", "\"Elegimos Express por tres razones concretas. Primero, es minimalista: no te obliga a seguir una estructura rígida, lo que nos dio libertad para organizar el proyecto en módulos según las funcionalidades del sistema. Segundo, es muy liviano y rápido, lo que es ideal para una aplicación de gestión de citas donde la respuesta debe ser inmediata. Tercero, tiene una comunidad enorme y documentación muy completa, lo que facilita resolver problemas durante el desarrollo. Alternativas como NestJS o Fastify también son válidas, pero Express es la opción más directa y probada para un proyecto de este tamaño.\"")
nota(doc, "PUNTO CLAVE: Express es el FRAMEWORK del backend. Node.js es el entorno que lo hace funcionar. Son cosas distintas — Express corre sobre Node.js.")
parrafo(doc, "Muestra app.js y explica:")
etiqueta(doc, "Explica app.js", "\"Este archivo es el punto de entrada del servidor. Aquí se configuran todas las rutas disponibles del sistema. Cada ruta le dice al servidor: cuando alguien pida esto, ejecuta este módulo. Por ejemplo, cuando alguien hace una petición a /api/appointments, el servidor sabe que debe ir al módulo de citas.\"")
parrafo(doc, "Muestra la carpeta modules/ y explica su estructura:")
parrafo(doc, "  • auth/ → \"Maneja el inicio de sesión y la seguridad. Verifica que el usuario exista y que la contraseña sea correcta.\"")
parrafo(doc, "  • appointments/ → \"Es el módulo principal. Aquí está toda la lógica para consultar horarios disponibles, reservar una cita y cancelarla.\"")
parrafo(doc, "  • notifications/ → \"Se encarga de enviar los correos electrónicos automáticos cuando se reserva una cita.\"")
parrafo(doc, "  • admin/ → \"Contiene las funciones exclusivas del administrador: bloquear y liberar horarios.\"")
parrafo(doc, "  • clients/ → \"Gestiona el registro básico del cliente al momento de hacer su reserva.\"")
etiqueta(doc, "Estructura de cada módulo", "\"Cada módulo tiene tres archivos con roles distintos. El archivo de rutas recibe la petición y la dirige. El controlador decide qué hacer con esa petición. El servicio ejecuta la lógica real: consultar la base de datos, guardar información, enviar el correo. Esta separación hace el código más ordenado y fácil de mantener.\"")
nota(doc, "PUNTO CLAVE: Esta estructura se llama arquitectura en capas. No tienes que decirlo así, pero sí explicar que cada archivo tiene una responsabilidad diferente.")

separador(doc)

# DATABASE
seccion(doc, "2.3 — BASE DE DATOS (donde se guarda todo)")
parrafo(doc, "Abre el archivo 'backend/prisma/schema.prisma' en VS Code.")
etiqueta(doc, "Qué decir", "\"La base de datos es donde se almacena toda la información del sistema: los usuarios, los clientes y las citas. Usamos PostgreSQL, que es un gestor de bases de datos relacional, ampliamente usado en proyectos profesionales. Para conectar el backend con la base de datos usamos Prisma, que es una herramienta que nos permite escribir las consultas de forma más clara y segura, sin escribir SQL directamente.\"")
parrafo(doc, "Muestra el schema.prisma y explica cada tabla:")
etiqueta(doc, "Tabla User", "\"Aquí se guardan los administradores del sistema. Tiene su nombre de usuario y su contraseña. La contraseña nunca se guarda tal como el usuario la escribe — se transforma con un algoritmo llamado bcrypt, que la convierte en un código irreconocible. Aunque alguien acceda a la base de datos, no puede saber cuál era la contraseña original.\"")
etiqueta(doc, "Tabla Client", "\"Aquí se registra la información básica del cliente que hace la reserva: su nombre y su correo electrónico. El correo es único, lo que evita registros duplicados.\"")
etiqueta(doc, "Tabla Appointment", "\"Esta es la tabla central del sistema. Cada fila representa un horario. Tiene la fecha y hora del turno, y un estado que puede ser AVAILABLE, es decir disponible, o RESERVED, es decir ocupado. Cuando un cliente reserva, el estado cambia de disponible a reservado, y se vincula con el cliente que lo tomó.\"")
nota(doc, "PUNTO IMPORTANTE sobre contraseñas: El instructor menciona que no deben guardarse en texto plano. Tú usas bcrypt, que es más seguro que MD5. Di exactamente esto: 'Las contraseñas se cifran con bcrypt antes de guardarse. Bcrypt es un algoritmo diseñado especialmente para proteger contraseñas, más robusto que otras alternativas.'")

doc.add_page_break()

# PARTE 3: DEMO
subtitulo(doc, "PARTE 3 — LA DEMO EN VIVO")
parrafo(doc, "Esta es la parte donde muestras el sistema funcionando. Practica este flujo varias veces antes de la sustentación para que salga sin tropiezos.")

seccion(doc, "Paso 1 — Abrir la aplicación")
etiqueta(doc, "Qué hacer", "Abre el navegador con la URL de la aplicación desplegada en Netlify.")
etiqueta(doc, "Qué decir", "\"Esta es la aplicación funcionando en producción. Está desplegada en internet y cualquier persona con el enlace puede acceder desde su celular o computador.\"")

seccion(doc, "Paso 2 — Vista del cliente")
etiqueta(doc, "Qué hacer", "Navega por la página principal como si fueras un cliente.")
etiqueta(doc, "Qué decir", "\"Esta es la vista que tiene el cliente. Puede ver los horarios disponibles del salón. Voy a seleccionar uno y registrar una cita en este momento.\"")

seccion(doc, "Paso 3 — Registrar una cita en vivo")
etiqueta(doc, "Qué hacer", "Selecciona un horario disponible, ingresa un nombre y correo de prueba, y confirma la reserva.")
etiqueta(doc, "Qué decir", "\"Ingreso el nombre del cliente y el correo electrónico. Al confirmar, el sistema reserva ese horario, lo marca como ocupado y envía automáticamente un correo de confirmación a esa dirección.\"")
nota(doc, "Ten listos datos de prueba: nombre 'Cliente Prueba', correo que puedas verificar en el momento si los instructores lo piden.")

seccion(doc, "Paso 4 — Vista del administrador")
etiqueta(doc, "Qué hacer", "Inicia sesión en el panel de administrador.")
etiqueta(doc, "Qué decir", "\"Ahora entro como administrador. Aquí puedo ver todas las citas registradas, cuáles están disponibles y cuáles están reservadas. También puedo bloquear horarios para que no estén disponibles, por ejemplo si el salón va a estar cerrado ese día.\"")

seccion(doc, "Paso 5 — Cierre de la demo")
etiqueta(doc, "Qué decir", "\"Con esto concluye la demostración. El sistema está funcionando en un entorno real, con base de datos real y notificaciones activas. Quedamos atentos a las preguntas de los instructores.\"")

doc.add_page_break()

# PREGUNTAS FRECUENTES
subtitulo(doc, "POSIBLES PREGUNTAS Y CÓMO RESPONDERLAS")

seccion(doc, "¿Por qué usaron JavaScript puro en el frontend y no un framework como React o Angular?")
parrafo(doc, "\"Para este proyecto el alcance era claro y concreto: una interfaz funcional sin complejidad innecesaria. JavaScript puro nos dio control total sin agregar dependencias que no íbamos a aprovechar. El resultado es una aplicación liviana que carga rápido y funciona en cualquier navegador.\"")

seccion(doc, "¿Cómo se protegen las contraseñas?")
parrafo(doc, "\"Las contraseñas se cifran con bcrypt antes de guardarse en la base de datos. Bcrypt convierte la contraseña en un código que no puede revertirse. Así, aunque alguien acceda a la base de datos, no puede conocer la contraseña original.\"")

seccion(doc, "¿Cómo sabe el sistema si eres cliente o administrador?")
parrafo(doc, "\"Al iniciar sesión, el servidor genera un token JWT, que es como un pase de identificación temporal. Ese token se guarda en el navegador y se envía en cada petición. El servidor lee el token y sabe exactamente quién eres y qué permisos tienes.\"")

seccion(doc, "¿Por qué PostgreSQL y no MySQL?")
parrafo(doc, "\"PostgreSQL es un gestor de bases de datos robusto, de código abierto y ampliamente usado en entornos profesionales. Soporta bien relaciones entre tablas, que es exactamente lo que necesitábamos: citas relacionadas con clientes. Ambos son válidos; elegimos PostgreSQL por su estabilidad y la integración directa con Prisma.\"")

seccion(doc, "¿Qué pasa si el sistema falla o no hay internet?")
parrafo(doc, "\"El sistema está desplegado en servidores en la nube, lo que garantiza alta disponibilidad. Para el backend usamos un túnel de Cloudflare que mantiene la conexión activa. En un escenario de producción real, se complementaría con un sistema de respaldo.\"")

seccion(doc, "¿Qué tiene el trabajo futuro?")
parrafo(doc, "\"Identificamos cuatro funcionalidades para la versión 2.0: asignación de empleados por servicio, perfiles de clientes con historial, catálogo de servicios con precios, y sistema de reseñas. Todo está documentado y listo para el siguiente ciclo de desarrollo.\"")

doc.add_page_break()

# TIPS FINALES
subtitulo(doc, "CONSEJOS PARA EL DÍA DE LA SUSTENTACIÓN")
parrafo(doc, "1. Practica en voz alta al menos 3 veces completo antes del día.")
parrafo(doc, "2. Abre todo lo que vas a mostrar antes de que empiece la sustentación: el PPT, VS Code con las carpetas y la aplicación en el navegador.")
parrafo(doc, "3. Habla despacio. Los instructores necesitan tiempo para procesar lo que ven en pantalla.")
parrafo(doc, "4. Si no sabes una respuesta, di: 'Ese es un punto que podemos profundizar en la versión 2.0' o 'En este momento no tengo ese dato preciso, pero lo puedo verificar.'")
parrafo(doc, "5. Conecta siempre lo que muestras con el problema que resuelve. No digas solo 'aquí está el código' — di 'aquí está el código que resuelve el problema del ausentismo.'")
parrafo(doc, "6. El módulo más importante que debes dominar es appointments (citas). Es el corazón del sistema.")
nota(doc, "Lo más importante: conoces el proyecto porque lo construiste. Confía en eso.")

# Guardar
ruta = "/Users/andres/Documents/Mara Beauty Studio/Guia_Sustentacion_Mara_Beauty_Studio.docx"
doc.save(ruta)
print(f"Documento guardado en: {ruta}")
